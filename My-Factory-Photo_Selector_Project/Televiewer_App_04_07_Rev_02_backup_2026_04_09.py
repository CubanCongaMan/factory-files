#!/usr/bin/env python3
"""
================================================================================
    OPTICAL TELEVIEWER PHOTO REPORT GENERATOR
    Version 1.0.0
================================================================================

    Generate professional photo reports from downhole televiewer images.
    
    For each borehole:
    1. Upload a televiewer image
    2. Enter header information (or pull from database)
    3. Generate a complete Word report with cropped depth images

    Copyright (c) 2026 Star Engineering, Inc.
    Developed with Factory AI

================================================================================
"""

import os
import sys
import re
import math
from copy import deepcopy
from datetime import datetime
from pathlib import Path

# Check dependencies
def check_dependencies():
    """Check and report missing dependencies."""
    missing = []
    
    try:
        from PIL import Image
    except ImportError:
        missing.append("Pillow")
    
    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")
    
    try:
        import requests
    except ImportError:
        missing.append("requests")
    
    if missing:
        print("\n" + "="*60)
        print("  MISSING DEPENDENCIES")
        print("="*60)
        print(f"\n  Please install the following packages:\n")
        print(f"    pip install {' '.join(missing)}")
        print("\n" + "="*60)
        sys.exit(1)

check_dependencies()

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """Application configuration."""
    
    # App Info
    APP_NAME = "Optical Televiewer Photo Report Generator"
    VERSION = "1.0.0"
    COMPANY = "Star Engineering, Inc."
    
    # Supabase
    SUPABASE_URL = "https://yyoaqeiinwrfdpfpgswc.supabase.co"
    SUPABASE_TABLE = "nada1"
    SUPABASE_ID_COLUMN = "vb_id"
    
    # Image Processing (calibrated values)
    PIXELS_PER_FOOT = 360
    DEPTH_OFFSET = -74
    HEADER_HEIGHT = 110
    RIGHT_CROP_X = 685
    LABEL_OFFSET = 20
    IMAGE_HEIGHT_INCHES = 7.0
    # First page lock rule:
    # - Include horizontal-axis header strip only on first page.
    # - Start first depth window below shallow PVC-obstructed interval.
    # - Keep subsequent pages at full FEET_PER_PAGE intervals.
    FIRST_PAGE_START_FT = 0.5
    FIRST_PAGE_END_FT = 3.0
    PVC_OBSCURED_TOP_FT = 2.0
    FEET_PER_PAGE = 3.0
    # Default strict mode: no overlap so page count tracks depth/3.
    OVERLAP_FEET = 0.0
    DEPTH_PARSE_TOLERANCE = 0.01
    IMAGE_WIDTH_CM = 9.0
    IMAGE_MIN_WIDTH_CM = 8.0
    IMAGE_HEIGHT_CM = 17.0
    # Last-page notes placement lock:
    # predict whether notes can fit under a short final image on the same page,
    # otherwise keep notes on the dedicated final notes page.
    NOTES_FONT_SIZE_PT = 11
    NOTES_LINE_HEIGHT_FACTOR = 1.2
    NOTES_EXTRA_PADDING_INCHES = 0.35
    NOTES_FIT_SAFETY_INCHES = 0.20
    MIN_LAST_IMAGE_HEIGHT_INCHES = 2.0
    
    # Report Notes
    NOTES = """Notes:
(1) Camera type and recording settings are as follows:
Type: Digital Optical Borehole Televiewer (Mount Sopris)
Frame Rate: 20 HZ
Exposure: Medium
Lighting: 20
Resolution: 900
Logging Speed: 2.7-2.9 ft/min
(2) Photograph Not-to-Scale
(3) Borehole orientation presented in the above image is relative to the local magnetic field"""


# =============================================================================
# CONSOLE INTERFACE
# =============================================================================

class Console:
    """Console UI helper."""
    
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BOLD = '\033[1m'
    END = '\033[0m'
    
    @classmethod
    def clear(cls):
        os.system('cls' if os.name == 'nt' else 'clear')
    
    @classmethod
    def header(cls):
        cls.clear()
        print(f"""
{cls.CYAN}{'='*65}{cls.END}
{cls.BOLD}{cls.CYAN}  OPTICAL TELEVIEWER PHOTO REPORT GENERATOR{cls.END}
{cls.CYAN}  Version {Config.VERSION} | {Config.COMPANY}{cls.END}
{cls.CYAN}{'='*65}{cls.END}
""")
    
    @classmethod
    def section(cls, title):
        print(f"\n{cls.BOLD}{cls.CYAN}  {title}{cls.END}")
        print(f"  {'-'*50}")
    
    @classmethod
    def success(cls, msg):
        print(f"  {cls.GREEN}✓{cls.END} {msg}")
    
    @classmethod
    def error(cls, msg):
        print(f"  {cls.RED}✗{cls.END} {msg}")
    
    @classmethod
    def info(cls, msg):
        print(f"  {cls.CYAN}→{cls.END} {msg}")
    
    @classmethod
    def warning(cls, msg):
        print(f"  {cls.YELLOW}!{cls.END} {msg}")
    
    @classmethod
    def prompt(cls, msg, default=None):
        if default:
            result = input(f"  {cls.YELLOW}>{cls.END} {msg} [{default}]: ").strip()
            return result if result else default
        return input(f"  {cls.YELLOW}>{cls.END} {msg}: ").strip()
    
    @classmethod
    def prompt_file(cls, msg):
        while True:
            path = cls.prompt(msg)
            if not path:
                cls.error("File path required.")
                continue
            if not os.path.exists(path):
                cls.error(f"File not found: {path}")
                continue
            return path
    
    @classmethod
    def confirm(cls, msg, default=True):
        hint = "Y/n" if default else "y/N"
        result = cls.prompt(f"{msg} ({hint})", "").lower()
        if not result:
            return default
        return result in ('y', 'yes')
    
    @classmethod
    def progress(cls, current, total, label="Progress"):
        pct = int(current / total * 100)
        bar = '█' * (pct // 5) + '░' * (20 - pct // 5)
        print(f"\r  {label}: {bar} {pct}%", end='', flush=True)
        if current == total:
            print()
    
    @classmethod
    def wait(cls):
        input(f"\n  {cls.CYAN}Press Enter to continue...{cls.END}")


# =============================================================================
# IMAGE PROCESSOR
# =============================================================================

class ImageProcessor:
    """Handles televiewer image cropping."""
    
    def __init__(self):
        self.cfg = Config()
    
    def detect_depth(self, height):
        """Auto-detect max depth from image height."""
        return int((height - self.cfg.DEPTH_OFFSET) / self.cfg.PIXELS_PER_FOOT)
    
    def get_row(self, depth, with_label=False):
        """Convert depth to pixel row."""
        row = self.cfg.DEPTH_OFFSET + int(depth * self.cfg.PIXELS_PER_FOOT)
        if with_label:
            row -= self.cfg.LABEL_OFFSET
        return max(0, row)

    def _detect_tick_peaks(self, gray_img, x_start, x_end, y_start=0):
        """Count ruler tick peaks in a vertical strip of a grayscale image."""
        pix = gray_img.load()
        width, height = gray_img.size
        x0 = max(0, min(width - 1, x_start))
        x1 = max(x0 + 1, min(width, x_end))
        ys = max(0, min(height, y_start))

        scores = []
        for y in range(ys, height):
            dark = 0
            for x in range(x0, x1):
                if pix[x, y] < 70:
                    dark += 1
            scores.append(dark)

        peaks = []
        for i in range(1, len(scores) - 1):
            # Threshold tuned for thin black ruler ticks.
            if scores[i] >= 6 and scores[i] >= scores[i - 1] and scores[i] >= scores[i + 1]:
                y = ys + i
                if not peaks or (y - peaks[-1]) > 8:
                    peaks.append(y)
        return len(peaks)

    def _find_ruler_boundary_x(self, gray_img):
        """Find x position where white ruler panel transitions to log texture."""
        w, h = gray_img.size
        pix = gray_img.load()
        y0 = max(0, h // 4)
        y1 = min(h, (3 * h) // 4)
        if y1 <= y0:
            return w // 2

        means = []
        for x in range(w):
            s = 0
            n = 0
            for y in range(y0, y1):
                s += pix[x, y]
                n += 1
            means.append((s / n) if n else 255)

        # Transition is typically sharp from bright ruler to darker borehole.
        for x in range(10, w - 10):
            if means[x - 1] > 220 and means[x] < 200:
                return x
        return w // 2

    def count_combined_ruler_ticks(self, crop_img, y_start=0):
        """Count combined left+right ruler ticks in cropped image."""
        gray = crop_img.convert('L')
        w, _ = gray.size
        left = self._detect_tick_peaks(gray, 0, min(35, w), y_start=y_start)

        boundary_x = self._find_ruler_boundary_x(gray)
        right = self._detect_tick_peaks(
            gray,
            max(0, boundary_x - 35),
            min(w, boundary_x + 1),
            y_start=y_start,
        )
        return left + right
    
    def crop_image(self, image_path, output_dir):
        """
        Crop televiewer image into depth intervals.
        
        Returns:
            tuple: (list of image paths, max depth)
        """
        img = Image.open(image_path)
        width, height = img.size
        basename = Path(image_path).stem
        
        max_depth = self.detect_depth(height)
        os.makedirs(output_dir, exist_ok=True)

        # Ensure deterministic outputs per run: remove stale JPG crops from
        # previous runs in this output folder.
        for name in os.listdir(output_dir):
            if name.lower().endswith('.jpg'):
                try:
                    os.remove(os.path.join(output_dir, name))
                except OSError:
                    # Non-fatal: continue with fresh crop generation.
                    pass
        
        # First-page exception is intentional and locked:
        # include axis header + skip shallow PVC-obscured interval by starting
        # at FIRST_PAGE_START_FT.
        first_start = self.cfg.FIRST_PAGE_START_FT
        first_end = min(self.cfg.FIRST_PAGE_END_FT, float(max_depth))
        intervals = [(first_start, first_end)]

        # From page 2 onward enforce full 3-ft windows with small overlap.
        step_ft = max(0.1, self.cfg.FEET_PER_PAGE - self.cfg.OVERLAP_FEET)
        start = self.cfg.FIRST_PAGE_END_FT
        max_depth_f = float(max_depth)
        while (start + self.cfg.FEET_PER_PAGE) <= (max_depth_f + self.cfg.DEPTH_PARSE_TOLERANCE):
            intervals.append((start, start + self.cfg.FEET_PER_PAGE))
            start += step_ft

        # Guarantee final post-first-page window still covers at least 3 ft.
        # This may overlap more near bottom by design to preserve scale.
        if max_depth_f > self.cfg.FIRST_PAGE_END_FT:
            forced_start = max(self.cfg.FIRST_PAGE_END_FT, max_depth_f - self.cfg.FEET_PER_PAGE)
            forced_interval = (forced_start, max_depth_f)
            if forced_interval[1] - forced_interval[0] >= (self.cfg.FEET_PER_PAGE - self.cfg.DEPTH_PARSE_TOLERANCE):
                last = intervals[-1]
                if abs(last[0] - forced_interval[0]) > self.cfg.DEPTH_PARSE_TOLERANCE or abs(last[1] - forced_interval[1]) > self.cfg.DEPTH_PARSE_TOLERANCE:
                    intervals.append(forced_interval)
        
        output_files = []
        
        for i, (start, end) in enumerate(intervals):
            Console.progress(i + 1, len(intervals), "Cropping")
            
            if i == 0:
                # First page lock: header + reduced depth window by design.
                body_start = self.get_row(start)
                body_end = self.get_row(end)
                crop_height = self.cfg.HEADER_HEIGHT + (body_end - body_start)
                
                crop = Image.new('RGB', (self.cfg.RIGHT_CROP_X, crop_height), 'white')
                header = img.crop((0, 0, self.cfg.RIGHT_CROP_X, self.cfg.HEADER_HEIGHT))
                body = img.crop((0, body_start, self.cfg.RIGHT_CROP_X, body_end))
                crop.paste(header, (0, 0))
                crop.paste(body, (0, self.cfg.HEADER_HEIGHT))

                # Tick-based correction: enforce visible depth scale using ruler
                # ticks (combined left+right, 10 ticks per foot).
                expected_ticks = int(round((end - start) * 10.0))
                max_extra = int(self.cfg.PIXELS_PER_FOOT * 2.0)  # cap expansion
                added = 0
                while True:
                    observed = self.count_combined_ruler_ticks(crop, y_start=self.cfg.HEADER_HEIGHT)
                    if observed >= int(expected_ticks * 0.9):
                        break
                    if body_end >= height or added >= max_extra:
                        break
                    step = max(1, int(self.cfg.PIXELS_PER_FOOT * 0.1))
                    body_end = min(height, body_end + step)
                    added += step
                    crop_height = self.cfg.HEADER_HEIGHT + (body_end - body_start)
                    crop = Image.new('RGB', (self.cfg.RIGHT_CROP_X, crop_height), 'white')
                    header = img.crop((0, 0, self.cfg.RIGHT_CROP_X, self.cfg.HEADER_HEIGHT))
                    body = img.crop((0, body_start, self.cfg.RIGHT_CROP_X, body_end))
                    crop.paste(header, (0, 0))
                    crop.paste(body, (0, self.cfg.HEADER_HEIGHT))
            else:
                # Subsequent images start at depth label
                row_start = self.get_row(start, with_label=True)
                row_end = min(self.get_row(end, with_label=True), height)

                # Tick-based correction: ensure ~3-ft pages visibly contain the
                # expected ruler tick count.
                expected_ticks = int(round((end - start) * 10.0))
                max_extra = int(self.cfg.PIXELS_PER_FOOT * 2.0)
                added = 0
                while True:
                    crop = img.crop((0, row_start, self.cfg.RIGHT_CROP_X, row_end))
                    observed = self.count_combined_ruler_ticks(crop)
                    if observed >= int(expected_ticks * 0.9):
                        break
                    if row_end >= height or added >= max_extra:
                        break
                    step = max(1, int(self.cfg.PIXELS_PER_FOOT * 0.1))
                    row_end = min(height, row_end + step)
                    added += step
                crop = img.crop((0, row_start, self.cfg.RIGHT_CROP_X, row_end))
            
            # Generate filename
            filename = f"{basename}_{start:.1f}-{end:.1f}ft.jpg"
            filepath = os.path.join(output_dir, filename)
            crop.save(filepath, 'JPEG', quality=95)
            output_files.append(filepath)
        
        return output_files, max_depth


# =============================================================================
# DATABASE CLIENT
# =============================================================================

class Database:
    """Supabase database client."""
    
    def __init__(self, api_key):
        self.url = Config.SUPABASE_URL
        self.key = api_key
        self.last_error = ""
    
    def _headers(self):
        headers = {
            "apikey": self.key,
            "Content-Type": "application/json",
            "Prefer": "return=representation"
        }
        # New Supabase secret keys (sb_secret_...) are not JWTs and should not
        # be sent as Authorization Bearer tokens.
        if self.key.startswith("eyJ"):
            headers["Authorization"] = f"Bearer {self.key}"
        return headers

    def _normalize_id(self, value):
        """Normalize boring IDs for tolerant comparison."""
        raw = str(value or "").strip().upper()
        return ''.join(ch for ch in raw if ch.isalnum())
    
    def test(self):
        """Test connection."""
        try:
            r = requests.get(
                f"{self.url}/rest/v1/{Config.SUPABASE_TABLE}?select=*&limit=1",
                headers=self._headers(),
                timeout=10
            )
            if r.status_code == 200:
                self.last_error = ""
                return True
            self.last_error = f"HTTP {r.status_code}: {r.text[:200]}"
            return False
        except Exception as e:
            self.last_error = str(e)
            return False
    
    def get_log(self, vb_id):
        """Get log by boring ID."""
        id_col = Config.SUPABASE_ID_COLUMN
        vb_id = str(vb_id or "").strip()
        if not vb_id:
            self.last_error = "Empty boring ID."
            return None

        endpoint = f"{self.url}/rest/v1/{Config.SUPABASE_TABLE}"
        candidate_cols = [id_col]

        def _try_query(column, matcher):
            return requests.get(
                endpoint,
                headers=self._headers(),
                params={
                    "select": "*",
                    "limit": 1,
                    column: matcher,
                },
                timeout=10,
            )

        def _is_missing_column_error(resp):
            if resp is None or resp.status_code != 400:
                return False
            body = (resp.text or "").lower()
            return "42703" in body or "does not exist" in body

        r = None
        data = []

        # 1) Exact match across candidate ID columns.
        for col in candidate_cols:
            r = _try_query(col, f"eq.{vb_id}")
            if r.status_code != 200:
                if _is_missing_column_error(r):
                    continue
                self.last_error = f"Lookup failed (HTTP {r.status_code}): {r.text[:200]}"
                return None
            data = r.json() if r.status_code == 200 else []
            if data:
                break

        # 2) Case-insensitive wildcard match across candidate columns.
        if not data:
            pattern = vb_id.replace('*', '')
            for col in candidate_cols:
                r = _try_query(col, f"ilike.*{pattern}*")
                if r.status_code != 200:
                    if _is_missing_column_error(r):
                        continue
                    self.last_error = f"Lookup failed (HTTP {r.status_code}): {r.text[:200]}"
                    return None
                data = r.json() if r.status_code == 200 else []
                if data:
                    break

        # 3) Normalized fallback (remove dashes/spaces, uppercase) by scanning a
        # reasonable slice of rows if targeted filters miss.
        if not data:
            r = requests.get(
                endpoint,
                headers=self._headers(),
                params={
                    "select": "*",
                    "limit": 2000,
                },
                timeout=20,
            )
            if r.status_code != 200:
                self.last_error = f"Lookup failed (HTTP {r.status_code}): {r.text[:200]}"
                return None
            rows = r.json() if r.status_code == 200 else []
            want = self._normalize_id(vb_id)
            for row in rows:
                for col in candidate_cols:
                    if self._normalize_id(row.get(col)) == want:
                        data = [row]
                        break
                if data:
                    break

        self.last_error = ""
        if not data:
            return None
        row = data[0]
        if 'vb_id_txt' not in row and id_col in row:
            row['vb_id_txt'] = row[id_col]
        return row
    
    def list_logs(self):
        """List all logs."""
        id_col = Config.SUPABASE_ID_COLUMN
        r = requests.get(
            f"{self.url}/rest/v1/{Config.SUPABASE_TABLE}?select=*&order={id_col}",
            headers=self._headers()
        )
        rows = r.json() if r.status_code == 200 else []
        for row in rows:
            if 'vb_id_txt' not in row and id_col in row:
                row['vb_id_txt'] = row[id_col]
        return rows
    
    def save_log(self, data):
        """Insert or update log."""
        id_col = Config.SUPABASE_ID_COLUMN
        id_val = data.get(id_col) or data.get('vb_id_txt')
        payload = dict(data)
        if id_col != 'vb_id_txt':
            payload[id_col] = id_val
            payload.pop('vb_id_txt', None)

        # Check if exists
        existing = self.get_log(id_val)
        if existing:
            r = requests.patch(
                f"{self.url}/rest/v1/{Config.SUPABASE_TABLE}?{id_col}=eq.{id_val}",
                headers=self._headers(),
                json=payload
            )
        else:
            r = requests.post(
                f"{self.url}/rest/v1/{Config.SUPABASE_TABLE}",
                headers=self._headers(),
                json=payload
            )
        return r.status_code in (200, 201)


# =============================================================================
# REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    """Generates Word reports from template."""
    
    def __init__(self):
        self.cfg = Config()
    
    def format_date(self, date_str):
        """Format date as 'Month DD, YYYY'."""
        if not date_str:
            return ""
        try:
            dt = datetime.strptime(str(date_str)[:10], "%Y-%m-%d")
            return dt.strftime("%B %d, %Y")
        except:
            return str(date_str)

    def build_replacements(self, header_data):
        """Build placeholder replacement map from database/header fields."""
        replacements = {}

        # Include every field from DB/manual input exactly as provided.
        for key, val in (header_data or {}).items():
            replacements[str(key)] = '' if val is None else str(val)

        # Backward/forward compatibility aliases for common ID keys.
        if 'vb_id' in replacements and 'vb_id_txt' not in replacements:
            replacements['vb_id_txt'] = replacements['vb_id']
        if 'vb_id_txt' in replacements and 'vb_id' not in replacements:
            replacements['vb_id'] = replacements['vb_id_txt']
        if 'north' in replacements and 'north_txt' not in replacements:
            replacements['north_txt'] = replacements['north']
        if 'north_txt' in replacements and 'north' not in replacements:
            replacements['north'] = replacements['north_txt']
        if 'easti' in replacements and 'easti_txt' not in replacements:
            replacements['easti_txt'] = replacements['easti']
        if 'easti_txt' in replacements and 'easti' not in replacements:
            replacements['easti'] = replacements['easti_txt']
        if 'ct_' in replacements and 'ct_txt' not in replacements:
            replacements['ct_txt'] = replacements['ct_']
        if 'ct_txt' in replacements and 'ct_' not in replacements:
            replacements['ct_'] = replacements['ct_txt']

        # Date placeholders can appear as either <field> or <field>_date.
        for key, val in list(replacements.items()):
            lk = key.lower()
            if lk.endswith('_date'):
                replacements[f"{key}_date"] = self.format_date(val)

        # Legacy aliases retained for older templates.
        legacy_aliases = {
            'north_txt': ['northing', 'north'],
            'easti_txt': ['easting', 'east'],
            'stat_num': ['station', 'station_num'],
            'ground_elev_num': ['ground_elev', 'ground_elevation'],
            'column_panel_txt': ['column_panel'],
            'column_panel_joint_txt': ['column_panel_joint'],
            'ct_txt': ['ct', 'ct_number'],
            'drill_by_txt': ['drilled_by', 'drill_by'],
            'op_tv_logger': ['tv_logger', 'op_tv_logger_txt'],
            'op_tv_date': ['tv_date', 'op_tv_date_date'],
        }
        for old_key, candidates in legacy_aliases.items():
            if old_key in replacements and replacements[old_key]:
                continue
            for candidate in candidates:
                if candidate in replacements and replacements[candidate]:
                    replacements[old_key] = replacements[candidate]
                    break

        # Keep old formatted date placeholders working.
        if 'drill_date' in replacements:
            replacements['drill_date_date'] = self.format_date(replacements['drill_date'])
        if 'op_tv_date' in replacements:
            replacements['op_tv_date'] = self.format_date(replacements['op_tv_date'])

        return replacements

    def is_canonical_template(self, doc):
        """Return True if template is single-page canonical televiewer format."""
        if len(doc.tables) != 3:
            return False
        shapes = [(len(t.rows), len(t.columns)) for t in doc.tables]
        return shapes == [(2, 2), (5, 4), (1, 2)]

    def canonical_image_cell(self, table):
        """Return image target cell for canonical Table 3 (1x2)."""
        return table.rows[0].cells[1]

    def get_canonical_page_prototype(self, doc):
        """Capture one canonical page body prototype (without section props)."""
        body = doc._element.body
        children = list(body.iterchildren())
        sect = children[-1] if children and children[-1].tag.endswith('sectPr') else None
        page_proto = []
        for el in children:
            if el is sect:
                continue
            # Stop at first explicit page break marker to keep single-page proto.
            if el.tag.endswith('p') and 'w:type="page"' in el.xml:
                break
            page_proto.append(deepcopy(el))
        return page_proto

    def append_canonical_page(self, doc, page_proto):
        """Append one canonical page before section properties."""
        body = doc._element.body

        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        body.insert(len(body) - 1, p)
        for el in page_proto:
            body.insert(len(body) - 1, deepcopy(el))

    def collect_canonical_image_cells(self, doc):
        """Collect one image cell per canonical page in document order."""
        cells = []
        for table in doc.tables:
            if len(table.rows) == 1 and len(table.columns) == 2:
                cells.append(self.canonical_image_cell(table))
        return cells

    def _prevent_row_split(self, row):
        """Prevent a table row from splitting across pages."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        existing = trPr.find(qn('w:cantSplit'))
        if existing is None:
            trPr.append(OxmlElement('w:cantSplit'))

    def _compact_paragraph(self, para):
        """Minimize paragraph spacing to keep page content deterministic."""
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.keep_together = True

    def apply_containment_lock(self, doc):
        """Harden Word layout so each 3-table mask stays on one page.

        This minimizes runtime pagination drift by disabling row splits and
        removing non-essential paragraph spacing inside/around tables.
        """
        for table in doc.tables:
            # Keep table geometry stable instead of letting Word re-autofit.
            table.autofit = False

            for row in table.rows:
                self._prevent_row_split(row)
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._compact_paragraph(para)

        # Also compact standalone spacer paragraphs between tables/pages.
        for para in doc.paragraphs:
            self._compact_paragraph(para)
    
    def find_image_cells(self, doc):
        """Find image area table indices.

        Strict rule: image pages must be true televiewer image pages (3x4 layout
        with an empty image slot in the first row). This avoids accidentally
        treating header/layout tables as image targets, which can shift content
        and corrupt report formatting.
        """
        modern_indices = []
        legacy_indices = []
        for i, table in enumerate(doc.tables):
            all_text = " ".join((cell.text or "") for row in table.rows for cell in row.cells).upper()

            # New template: image page uses a 3x4 table with an empty first row,
            # and label text in lower rows. Keep this strict so page headers
            # remain in place and the 3-ft image presentation remains consistent.
            top_row_text = " ".join((cell.text or "").strip() for cell in table.rows[0].cells)
            if len(table.rows) == 3 and len(table.columns) == 4 and "OPTICAL TELEVIEWER IMAGE LOG" in all_text and not top_row_text.strip():
                modern_indices.append(i)

            # Legacy template: image page as empty 1x1 table. Keep only if not
            # the final trailing notes table.
            if len(table.rows) == 1 and len(table.columns) == 1:
                if not table.rows[0].cells[0].text.strip() and i != len(doc.tables) - 1:
                    legacy_indices.append(i)

        return list(dict.fromkeys(modern_indices)), list(dict.fromkeys(legacy_indices))

    def select_consistent_image_cells(self, doc, required_images=0):
        """Select a single, consistent image-page layout family.

        Hard rule: do not mix modern (3x4) and legacy (1x1) image page layouts
        in one generated report, because it causes page-area geometry drift.
        """
        modern_indices, legacy_indices = self.find_image_cells(doc)

        if modern_indices and legacy_indices:
            # Prefer modern pages when they alone satisfy capacity.
            if required_images and required_images <= len(modern_indices):
                return modern_indices, "modern", len(modern_indices), len(legacy_indices)
            # If both are present and modern is insufficient, block generation
            # to preserve strict page-size consistency.
            raise ValueError(
                "Template mixes modern (3x4) and legacy (1x1) image page layouts. "
                f"Consistent modern capacity is {len(modern_indices)} pages, "
                f"but {required_images} image pages are required. "
                "Use a template where all image pages share one layout family."
            )

        if modern_indices:
            return modern_indices, "modern", len(modern_indices), 0
        return legacy_indices, "legacy", 0, len(legacy_indices)

    def parse_depth_interval(self, image_path):
        """Parse depth interval from filename like *_3.0-6ft.jpg."""
        name = os.path.basename(image_path)
        match = re.search(r"_(\d+(?:\.\d+)?)-(\d+(?:\.\d+)?)ft", name, re.IGNORECASE)
        if not match:
            return None
        return float(match.group(1)), float(match.group(2))

    def target_image_height_cm(self, image_path):
        """Return display height in cm using locked printed scale.

        User rule: each 0.1 ft tick should print at 0.5 cm spacing.
        Therefore 1.0 ft = 5.0 cm and a 3-ft interval = 15.0 cm.
        """
        interval = self.parse_depth_interval(image_path)
        if not interval:
            return 15.0
        start, end = interval
        span_ft = max(0.0, end - start)
        return span_ft * 5.0

    def validate_three_foot_geometry(self, images):
        """Ensure full 3-ft pages keep identical pixel geometry.

        This enforces the strict reporting rule: when a page represents exactly
        3.0 feet of log, the rendered image geometry must be consistent.
        """
        baseline_size = None
        offenders = []
        for image_path in images:
            interval = self.parse_depth_interval(image_path)
            if not interval:
                continue
            start, end = interval
            span = end - start
            if abs(span - self.cfg.FEET_PER_PAGE) > self.cfg.DEPTH_PARSE_TOLERANCE:
                continue

            with Image.open(image_path) as img:
                size = img.size

            if baseline_size is None:
                baseline_size = size
                continue

            if size != baseline_size:
                offenders.append((os.path.basename(image_path), size, baseline_size))

        if offenders:
            details = "; ".join(
                f"{name}: {size[0]}x{size[1]} expected {expected[0]}x{expected[1]}"
                for name, size, expected in offenders[:5]
            )
            raise ValueError(
                "3-ft page geometry mismatch detected. "
                "To preserve vertical scale and consistent image width, all full "
                f"3-ft crops must share identical pixel size. {details}"
            )

    def estimate_notes_height_inches(self):
        """Estimate rendered height of locked notes block in inches."""
        line_count = max(1, len(self.cfg.NOTES.splitlines()))
        line_height_in = (self.cfg.NOTES_FONT_SIZE_PT * self.cfg.NOTES_LINE_HEIGHT_FACTOR) / 72.0
        return (line_count * line_height_in) + self.cfg.NOTES_EXTRA_PADDING_INCHES

    def predict_notes_same_page(self, last_image_path):
        """Predict if notes fit below the final short-depth image on same page.

        This rule avoids creating mostly blank pages for short final intervals
        while keeping the notes block intact (never split across pages).
        """
        interval = self.parse_depth_interval(last_image_path)
        if not interval:
            return False, self.cfg.IMAGE_HEIGHT_INCHES

        start, end = interval
        span = max(0.0, end - start)
        # Only short final segments are candidates for same-page notes.
        if span >= (self.cfg.FEET_PER_PAGE - self.cfg.DEPTH_PARSE_TOLERANCE):
            return False, self.cfg.IMAGE_HEIGHT_INCHES

        proportional_height = self.cfg.IMAGE_HEIGHT_INCHES * (span / self.cfg.FEET_PER_PAGE)
        last_image_height = max(self.cfg.MIN_LAST_IMAGE_HEIGHT_INCHES, proportional_height)
        available_below_image = self.cfg.IMAGE_HEIGHT_INCHES - last_image_height
        required = self.estimate_notes_height_inches() + self.cfg.NOTES_FIT_SAFETY_INCHES
        return available_below_image >= required, last_image_height

    def find_notes_table(self, doc):
        """Find the notes table index, if present."""
        for i in range(len(doc.tables) - 1, -1, -1):
            table = doc.tables[i]
            if len(table.rows) == 1 and len(table.columns) == 1 and not table.rows[0].cells[0].text.strip():
                return i
        return None

    def get_image_slot_count(self, template_path):
        """Return number of image placeholders available in a template."""
        doc = Document(template_path)
        indices, _, _, _ = self.select_consistent_image_cells(doc, required_images=0)
        return len(indices)
    
    def generate(self, template_path, images, header_data, output_path):
        """
        Generate the photo report.
        
        Args:
            template_path: Path to Word template
            images: List of cropped image paths
            header_data: Dictionary with header field values
            output_path: Output file path
        """
        doc = Document(template_path)
        canonical_mode = self.is_canonical_template(doc)

        if canonical_mode:
            if not images:
                raise ValueError("No cropped images generated for report.")

            # Build N canonical pages from one canonical page template.
            page_proto = self.get_canonical_page_prototype(doc)
            for _ in range(max(0, len(images) - 1)):
                self.append_canonical_page(doc, page_proto)
            image_cells = self.collect_canonical_image_cells(doc)
            if len(image_cells) < len(images):
                raise ValueError(
                    f"Canonical page clone mismatch: {len(image_cells)} image cells for {len(images)} images."
                )

            replacements = self.build_replacements(header_data)

            total_steps = len(images) + 2
            step = 0

            # Insert images with strict placement rules for Table 3 area.
            for i, (cell, img_path) in enumerate(zip(image_cells, images)):
                step += 1
                Console.progress(step, total_steps, "Building report")

                for para in cell.paragraphs:
                    para._element.getparent().remove(para._element)

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                width_cm = max(self.cfg.IMAGE_MIN_WIDTH_CM, self.cfg.IMAGE_WIDTH_CM)
                height_cm = self.target_image_height_cm(img_path)
                para.add_run().add_picture(img_path, width=Cm(width_cm), height=Cm(height_cm))

            # Fill green placeholders on every cloned page from one in-memory
            # header_data record to avoid repeated DB calls.
            step += 1
            Console.progress(step, total_steps, "Building report")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            changed = False
                            for key in sorted(replacements.keys(), key=len, reverse=True):
                                val = replacements[key]
                                if key in text:
                                    text = text.replace(key, str(val) if val else "")
                                    changed = True
                            if changed:
                                if para.runs:
                                    for run in para.runs:
                                        run.text = ''
                                    para.runs[0].text = text
                                else:
                                    para.add_run(text)

            # Last-page notes lock: place notes on same page if predicted to fit,
            # otherwise create one extra canonical page for notes only.
            same_page_notes, _ = self.predict_notes_same_page(images[-1])
            if same_page_notes:
                notes_para = image_cells[-1].add_paragraph()
                notes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = notes_para.add_run(self.cfg.NOTES)
                run.font.size = Pt(self.cfg.NOTES_FONT_SIZE_PT)
                run.font.name = 'Arial'
            else:
                self.append_canonical_page(doc, page_proto)
                note_cell = self.collect_canonical_image_cells(doc)[-1]
                for para in note_cell.paragraphs:
                    para._element.getparent().remove(para._element)
                note_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                notes_para = note_cell.add_paragraph()
                notes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = notes_para.add_run(self.cfg.NOTES)
                run.font.size = Pt(self.cfg.NOTES_FONT_SIZE_PT)
                run.font.name = 'Arial'

            # Containment lock: keep each cloned 3-table mask in one page.
            self.apply_containment_lock(doc)

            step += 1
            Console.progress(step, total_steps, "Building report")
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
            doc.save(output_path)
            return output_path

        image_indices, slot_family, modern_count, legacy_count = self.select_consistent_image_cells(
            doc, required_images=len(images)
        )

        # Enforce strict consistency for full 3-ft page crops before inserting
        # into the report template.
        self.validate_three_foot_geometry(images)

        if not image_indices:
            raise ValueError("No image slots found in template. Use a compatible televiewer template.")
        if len(images) > len(image_indices):
            raise ValueError(
                f"Template capacity exceeded: {len(images)} images but template has {len(image_indices)} image slots."
            )
        
        replacements = self.build_replacements(header_data)

        notes_idx = self.find_notes_table(doc)
        same_page_notes = False
        last_image_height_inches = self.cfg.IMAGE_HEIGHT_INCHES
        if images:
            # Predict whether to keep notes on last image page or dedicated page.
            same_page_notes, last_image_height_inches = self.predict_notes_same_page(images[-1])

        if same_page_notes and notes_idx is None:
            # No notes container exists; fallback safely to dedicated-page mode.
            same_page_notes = False
        
        total_steps = len(images) + 2
        step = 0
        
        # Insert images
        last_image_cell = None
        for i, (idx, img_path) in enumerate(zip(image_indices[:len(images)], images)):
            step += 1
            Console.progress(step, total_steps, "Building report")
            
            table = doc.tables[idx]
            cell = table.rows[0].cells[0]
            
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            target_height = self.target_image_height_cm(img_path) / 2.54
            if same_page_notes and i == (len(images) - 1):
                # Soft rule: shrink only the short final segment to free room
                # for notes on the same page while preserving width.
                target_height = last_image_height_inches
            para.add_run().add_picture(img_path, height=Inches(target_height))

            if i == (len(images) - 1):
                last_image_cell = cell
        
        # Replace placeholders
        step += 1
        Console.progress(step, total_steps, "Building report")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        changed = False
                        # Replace longer keys first to avoid substring collisions
                        # (e.g., replacing "easti" inside "easti_txt").
                        for key in sorted(replacements.keys(), key=len, reverse=True):
                            val = replacements[key]
                            if key in text:
                                text = text.replace(key, str(val) if val else "")
                                changed = True
                        if changed:
                            if para.runs:
                                for run in para.runs:
                                    run.text = ''
                                para.runs[0].text = text
                            else:
                                para.add_run(text)
        
        # Add notes with lock behavior:
        # - Same-page mode for short final segments when notes fully fit.
        # - Otherwise dedicated final notes page.
        if same_page_notes and last_image_cell is not None:
            notes_para = last_image_cell.add_paragraph()
            run = notes_para.add_run(self.cfg.NOTES)
            run.font.size = Pt(self.cfg.NOTES_FONT_SIZE_PT)
            run.font.name = 'Arial'

            # Remove dedicated notes table to avoid a mostly blank extra page.
            if notes_idx is not None:
                notes_table = doc.tables[notes_idx]
                notes_table._element.getparent().remove(notes_table._element)
        elif notes_idx is not None:
            table = doc.tables[notes_idx]
            cell = table.rows[0].cells[0]
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            notes_para = cell.add_paragraph()
            run = notes_para.add_run(self.cfg.NOTES)
            run.font.size = Pt(self.cfg.NOTES_FONT_SIZE_PT)
            run.font.name = 'Arial'

        # Containment lock for non-canonical templates as well.
        self.apply_containment_lock(doc)
        
        # Save
        step += 1
        Console.progress(step, total_steps, "Building report")
        
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        
        return output_path


# =============================================================================
# MAIN APPLICATION
# =============================================================================

class App:
    """Main application."""
    
    def __init__(self):
        self.db = None
        self.api_key = None

    @staticmethod
    def sanitize_filename_component(value):
        """Return a Windows-safe filename component."""
        text = str(value or "Unknown").strip()
        # Replace Windows-invalid filename characters.
        text = re.sub(r'[<>:"/\\|?*]', '_', text)
        # Windows disallows trailing dots/spaces in file or folder names.
        text = text.rstrip(' .')
        return text or "Unknown"

    @staticmethod
    def next_report_revision_path(output_dir, vb_id):
        """Return next available revisioned report path.

        Example output name:
        Optical_Televiewer_Report_H2-VB-12_Rev_03.docx
        """
        base = f"Optical_Televiewer_Report_{vb_id}"
        pattern = re.compile(rf"^{re.escape(base)}(?:_Rev_(\d{{2}}))?\.docx$", re.IGNORECASE)

        max_rev = 0
        for name in os.listdir(output_dir):
            match = pattern.match(name)
            if not match:
                continue

            rev_text = match.group(1)
            if rev_text is None:
                # Legacy non-revision file counts as Rev_00 baseline.
                max_rev = max(max_rev, 0)
            else:
                try:
                    max_rev = max(max_rev, int(rev_text))
                except ValueError:
                    pass

        next_rev = max_rev + 1
        filename = f"{base}_Rev_{next_rev:02d}.docx"
        return os.path.join(output_dir, filename)
    
    def run(self):
        """Run the application."""
        # Disable colors on Windows if needed
        if os.name == 'nt':
            try:
                import ctypes
                ctypes.windll.kernel32.SetConsoleMode(
                    ctypes.windll.kernel32.GetStdHandle(-11), 7)
            except:
                Console.CYAN = Console.GREEN = Console.YELLOW = ''
                Console.RED = Console.BOLD = Console.END = ''
        
        self.main_menu()
    
    def main_menu(self):
        """Show main menu."""
        while True:
            Console.header()
            Console.section("MAIN MENU")
            print("""
    [1] Generate Photo Report (with database)
    [2] Generate Photo Report (manual entry)
    [3] Crop Images Only
    [4] View Database Logs
    [Q] Quit
""")
            choice = Console.prompt("Select option").upper()
            
            if choice == '1':
                self.generate_with_database()
            elif choice == '2':
                self.generate_manual()
            elif choice == '3':
                self.crop_only()
            elif choice == '4':
                self.view_logs()
            elif choice == 'Q':
                Console.info("Goodbye!")
                print()
                break
    
    def connect_db(self):
        """Ensure database connection."""
        if self.db:
            return True
        
        Console.section("DATABASE CONNECTION")
        self.api_key = Console.prompt("Supabase service_role key (sb_secret_... or eyJ...)").strip()
        
        if not self.api_key:
            Console.error("API key required.")
            return False

        # Guardrail: users sometimes enter boring IDs at this prompt.
        if not (self.api_key.startswith("sb_secret_") or self.api_key.startswith("eyJ")):
            Console.error("That does not look like a Supabase API key.")
            Console.info("Enter the project API key (starts with sb_secret_... or eyJ...).")
            return False
        
        self.db = Database(self.api_key)
        
        Console.info("Testing connection...")
        if not self.db.test():
            Console.error("Connection failed.")
            if self.db.last_error:
                Console.warning(self.db.last_error)
            Console.info(f"Supabase URL: {Config.SUPABASE_URL}")
            Console.info("Use a key from this same Supabase project URL.")
            self.db = None
            return False
        
        Console.success("Connected to database!")
        return True
    
    def generate_with_database(self):
        """Generate report using database for header data."""
        Console.header()
        Console.section("GENERATE PHOTO REPORT (Database Mode)")

        # Get boring ID
        vb_id = Console.prompt("Verification Boring ID (e.g., H7-VB-04)").strip()
        if not vb_id:
            Console.error("Boring ID required.")
            Console.wait()
            return

        if not self.connect_db():
            Console.wait()
            return
        
        # Fetch from database
        Console.info(f"Looking up {vb_id}...")
        header_data = self.db.get_log(vb_id)
        
        if not header_data:
            Console.warning(f"No data found for {vb_id}")
            if self.db.last_error:
                Console.warning(self.db.last_error)
            if Console.confirm("Enter header data manually?"):
                header_data = self.get_header_input(vb_id)
                if Console.confirm("Save to database?"):
                    if self.db.save_log(header_data):
                        Console.success("Saved to database!")
                    else:
                        Console.error("Failed to save.")
            else:
                Console.wait()
                return
        else:
            Console.success(f"Found: {header_data['vb_id_txt']}")
        
        # Get files
        self.generate_report(header_data)
    
    def generate_manual(self):
        """Generate report with manual header entry."""
        Console.header()
        Console.section("GENERATE PHOTO REPORT (Manual Entry)")
        
        vb_id = Console.prompt("Verification Boring ID")
        if not vb_id:
            Console.error("Boring ID required.")
            Console.wait()
            return
        
        header_data = self.get_header_input(vb_id)
        self.generate_report(header_data)
    
    def get_header_input(self, vb_id):
        """Get header data from user input."""
        Console.section("ENTER HEADER DATA")
        
        return {
            'vb_id': vb_id,
            'north': Console.prompt("Northing"),
            'easti': Console.prompt("Easting"),
            'stat_num': Console.prompt("Station"),
            'ground_elev_num': Console.prompt("Ground Elevation", "0"),
            'column_panel_txt': Console.prompt("Column/Panel"),
            'column_panel_joint_txt': Console.prompt("Column/Panel Joint"),
            'ct_': Console.prompt("CT Number"),
            'drill_date': Console.prompt("Drill Date (YYYY-MM-DD)"),
            'drill_by_txt': Console.prompt("Drilled By"),
            'op_tv_logger': Console.prompt("TV Logger"),
            'op_tv_date': Console.prompt("TV Date (YYYY-MM-DD)")
        }
    
    def generate_report(self, header_data):
        """Generate the photo report."""
        Console.section("FILE SELECTION")
        
        # Get source image
        image_path = Console.prompt_file("Televiewer image file (JPG)")
        
        # Get template
        template_default = "Televiewer_Log_Template_A_one__Page_Example_with_No_Image_Rev_00.docx"
        if os.path.exists(template_default):
            template_path = Console.prompt(f"Word template", template_default)
            if not template_path:
                template_path = template_default
        else:
            template_path = Console.prompt_file("Word template file")
        
        # Output directory
        output_dir = Console.prompt("Output directory", "./output")
        os.makedirs(output_dir, exist_ok=True)
        
        # Process
        Console.section("GENERATING REPORT")
        
        try:
            # Crop images
            Console.info("Cropping televiewer image...")
            processor = ImageProcessor()
            cropped_dir = os.path.join(output_dir, "cropped_images")
            images, max_depth = processor.crop_image(image_path, cropped_dir)
            Console.success(f"Created {len(images)} images (depth: {max_depth} ft)")

            # Quick sanity check requested by user:
            # if image-page count materially exceeds depth/3 expectation,
            # report likely has unnecessary extra pages.
            expected_min_pages = max(
                1,
                math.ceil((max_depth - Config.FIRST_PAGE_START_FT) / Config.FEET_PER_PAGE)
            )
            if len(images) > expected_min_pages:
                raise ValueError(
                    "Sanity check failed: generated image pages exceed depth/3 expectation. "
                    f"Generated image pages: {len(images)} | Expected minimum: {expected_min_pages}. "
                    "Check overlap and depth calibration settings."
                )
            
            # Generate report
            Console.info("Building Word report...")
            generator = ReportGenerator()
            raw_vb_id = header_data.get('vb_id_txt') or header_data.get('vb_id') or "Unknown"
            vb_id = self.sanitize_filename_component(raw_vb_id)
            required_pages = max(1, (int(max_depth) + 2) // 3)

            template_doc = Document(template_path)
            if generator.is_canonical_template(template_doc):
                slot_count = len(images)
            else:
                slot_count = generator.get_image_slot_count(template_path)

            if slot_count <= 0:
                raise ValueError("Selected template has no compatible image slots.")

            if len(images) > slot_count:
                short_by = len(images) - slot_count
                recommended = required_pages + 2
                Console.error("You are going to need a bigger boat (template).")
                Console.warning(
                    f"Template capacity: {slot_count} image pages | Needed: {len(images)} image pages."
                )
                Console.info(f"Borehole depth: {max_depth} ft -> minimum pages at 3 ft/page: {required_pages}.")
                Console.info(f"Short by {short_by} pages. Recommended template size: {recommended}+ pages.")
                Console.info("Load a larger template and run again.")
                Console.wait()
                return

            output_path = self.next_report_revision_path(output_dir, vb_id)
            generator.generate(template_path, images, header_data, output_path)

            Console.success("Report generated successfully!")
            print(f"\n  Output: {output_path}")
            print(f"  Images: {cropped_dir}")
            
        except Exception as e:
            Console.error(f"Failed: {str(e)}")
        
        Console.wait()
    
    def crop_only(self):
        """Crop images without generating report."""
        Console.header()
        Console.section("CROP IMAGES ONLY")
        
        image_path = Console.prompt_file("Televiewer image file (JPG)")
        output_dir = Console.prompt("Output directory", "./cropped_images")
        
        Console.section("PROCESSING")
        
        try:
            processor = ImageProcessor()
            images, max_depth = processor.crop_image(image_path, output_dir)
            Console.success(f"Created {len(images)} images (depth: {max_depth} ft)")
            print(f"\n  Output: {output_dir}")
        except Exception as e:
            Console.error(f"Failed: {str(e)}")
        
        Console.wait()
    
    def view_logs(self):
        """View database logs."""
        Console.header()
        Console.section("DATABASE LOGS")
        
        if not self.connect_db():
            Console.wait()
            return
        
        logs = self.db.list_logs()
        
        if not logs:
            Console.warning("No logs found in database.")
        else:
            print(f"\n  {'Boring ID':<15} {'Drill Date':<15}")
            print(f"  {'-'*30}")
            for log in logs:
                drill = str(log.get('drill_date', ''))[:10]
                boring_id = str(log.get('vb_id_txt') or log.get(Config.SUPABASE_ID_COLUMN) or '')
                print(f"  {boring_id:<15} {drill:<15}")
            print(f"\n  Total: {len(logs)} logs")
        
        Console.wait()


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    try:
        App().run()
    except KeyboardInterrupt:
        print("\n\n  Interrupted. Goodbye!")
        sys.exit(0)
