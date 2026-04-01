#!/usr/bin/env python3
# 04/01/2026 Photo Selector Version 26 Rev 04
#Fix: New template auto-selection based on filename patterns (FRONT/BACK, WV/CU)
#Fix: Dynamic preview sizing - fills available space
#Author: RLS & Factory AI Droids

import os, shutil, re, tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk

try:
    import ttkbootstrap as ttk
    from ttkbootstrap.constants import *
    THEME_AVAILABLE = True
except ImportError:
    from tkinter import ttk
    THEME_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docxcompose.composer import Composer
    DOCX_OK = True
except:
    DOCX_OK = False

class App:
    def __init__(self, root):
        self.root = root
        root.title("Photo Selector v26")
        root.geometry("1400x900")
        
        self.bg_color = "#1a2942"
        self.fg_color = "#ffffff"
        self.accent_color = "#3d5a80"
        self.highlight_color = "#4a90d9"
        
        self.photos, self.thumbs, self.checks = [], {}, {}
        self.preview_img, self.original_image, self.current_photo_path = None, None, None
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.output_folder, self.renamed_photos = "", []
        self.crop_mode, self.crop_coords, self.crop_start, self.active_handle = False, [50,50,200,200], None, None
        self.vb_id = tk.StringVar()
        self.path = tk.StringVar()
        
        self.sorted_images = []
        self.preview_checks = {}
        self.preview_thumbs = {}
        self.current_page = 0
        self.pages = []
        
        self.templates_folder = tk.StringVar()
        self.templates = {}
        self.page_templates = {}
        
        self.style = ttk.Style()
        if not THEME_AVAILABLE:
            self.configure_fallback_style()
        
        main_container = tk.Frame(root, bg=self.bg_color)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        header = tk.Frame(main_container, bg=self.accent_color, height=50)
        header.pack(fill=tk.X, padx=0, pady=0)
        header.pack_propagate(False)
        
        title_label = tk.Label(header, text="📷 PHOTO SELECTOR & REPORT GENERATOR", 
                               font=('Arial', 16, 'bold'), bg=self.accent_color, fg=self.fg_color)
        title_label.pack(side=tk.LEFT, padx=20, pady=10)
        
        version_label = tk.Label(header, text="v26.04", font=('Arial', 10), 
                                  bg=self.accent_color, fg="#aabbcc")
        version_label.pack(side=tk.RIGHT, padx=20, pady=15)
        
        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.tab1 = tk.Frame(self.notebook, bg=self.bg_color)
        self.notebook.add(self.tab1, text="  📁 Photo Selector  ")
        self.build_photo_selector_tab()
        
        self.tab2 = tk.Frame(self.notebook, bg=self.bg_color)
        self.notebook.add(self.tab2, text="  📄 Report Preview  ")
        self.build_report_preview_tab()
        
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_change)

    def configure_fallback_style(self):
        self.style.configure('TNotebook', background=self.bg_color)
        self.style.configure('TNotebook.Tab', padding=[20, 10], font=('Arial', 10, 'bold'))
        self.style.map('TNotebook.Tab', 
                       background=[('selected', self.highlight_color), ('!selected', self.accent_color)],
                       foreground=[('selected', self.fg_color), ('!selected', '#cccccc')])

    def create_styled_label(self, parent, text, font_size=10, bold=False):
        font_weight = 'bold' if bold else 'normal'
        return tk.Label(parent, text=text, font=('Arial', font_size, font_weight),
                       bg=self.bg_color, fg=self.fg_color)

    def create_styled_button(self, parent, text, command, width=12):
        btn = tk.Button(parent, text=text, command=command, width=width,
                       font=('Arial', 9, 'bold'), bg=self.accent_color, fg=self.fg_color,
                       activebackground=self.highlight_color, activeforeground=self.fg_color,
                       relief=tk.FLAT, cursor='hand2', padx=10, pady=5)
        btn.bind('<Enter>', lambda e: btn.config(bg=self.highlight_color))
        btn.bind('<Leave>', lambda e: btn.config(bg=self.accent_color))
        return btn

    def create_styled_entry(self, parent, textvariable, width=40):
        entry = tk.Entry(parent, textvariable=textvariable, width=width,
                        font=('Arial', 10), bg='#2a3f5f', fg=self.fg_color,
                        insertbackground=self.fg_color, relief=tk.FLAT,
                        highlightthickness=2, highlightcolor=self.highlight_color,
                        highlightbackground=self.accent_color)
        return entry

    def create_styled_labelframe(self, parent, text):
        frame = tk.LabelFrame(parent, text=text, font=('Arial', 10, 'bold'),
                             bg=self.bg_color, fg=self.fg_color, padx=10, pady=10,
                             relief=tk.GROOVE, bd=2)
        return frame

    def build_photo_selector_tab(self):
        top = tk.Frame(self.tab1, bg=self.bg_color, pady=10)
        top.pack(fill=tk.X, padx=10)
        
        self.create_styled_label(top, "Boring ID:", 10, True).pack(side=tk.LEFT)
        self.create_styled_entry(top, self.vb_id, 15).pack(side=tk.LEFT, padx=(5, 20))
        
        self.create_styled_label(top, "Folder:", 10, True).pack(side=tk.LEFT)
        self.create_styled_entry(top, self.path, 35).pack(side=tk.LEFT, padx=5)
        
        self.create_styled_button(top, "Browse", self.browse, 8).pack(side=tk.LEFT, padx=3)
        self.create_styled_button(top, "Load", self.load, 8).pack(side=tk.LEFT, padx=3)
        self.create_styled_button(top, "Select All", self.sel_all, 10).pack(side=tk.LEFT, padx=(15,3))
        self.create_styled_button(top, "Clear All", self.clr_all, 10).pack(side=tk.LEFT, padx=3)
        
        main = tk.PanedWindow(self.tab1, orient=tk.HORIZONTAL, bg=self.bg_color, sashwidth=5, sashrelief=tk.RAISED)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        left = self.create_styled_labelframe(main, "📷 Photos - Click to Preview")
        
        self.canvas = tk.Canvas(left, bg='#243554', highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        self.inner = tk.Frame(self.canvas, bg='#243554')
        self.canvas.create_window((0,0), window=self.inner, anchor=tk.NW)
        
        right = tk.Frame(main, bg=self.bg_color)
        
        main.add(left, minsize=300)
        main.add(right, minsize=400)
        self.root.update_idletasks()
        main.sash_place(0, 800, 0)
        
        pf = self.create_styled_labelframe(right, "🔍 Preview")
        pf.pack(fill=tk.BOTH, expand=True)
        
        self.pcv = tk.Canvas(pf, bg='#243554', width=520, height=350, highlightthickness=0)
        self.pcv.pack(fill=tk.BOTH, expand=True)
        
        zf = tk.Frame(pf, bg=self.bg_color)
        zf.pack(fill=tk.X, pady=5)
        self.create_styled_button(zf, "Zoom+", self.zoom_in, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(zf, "Zoom-", self.zoom_out, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(zf, "Reset", self.zoom_reset, 7).pack(side=tk.LEFT, padx=2)
        self.crop_btn = self.create_styled_button(zf, "Crop", self.toggle_crop, 7)
        self.crop_btn.pack(side=tk.LEFT, padx=10)
        
        self.crop_frame = tk.Frame(pf, bg=self.bg_color)
        self.create_styled_button(self.crop_frame, "Apply", self.apply_crop, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(self.crop_frame, "Cancel", self.cancel_crop, 7).pack(side=tk.LEFT, padx=2)
        
        pnf = tk.Frame(pf, bg=self.bg_color)
        pnf.pack(fill=tk.X, pady=5)
        self.create_styled_button(pnf, "▲", lambda:self.pan(0,-30), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "▼", lambda:self.pan(0,30), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "◀", lambda:self.pan(-30,0), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "▶", lambda:self.pan(30,0), 4).pack(side=tk.LEFT, padx=2)
        
        self.lbl_n = self.create_styled_label(pf, "Name: -", 9)
        self.lbl_n.pack(anchor=tk.W, pady=2)
        self.lbl_s = self.create_styled_label(pf, "Size: -", 9)
        self.lbl_s.pack(anchor=tk.W, pady=2)
        
        sf = self.create_styled_labelframe(right, "✓ Selected Photos")
        sf.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        self.slist = tk.Listbox(sf, height=5, font=('Arial', 9),
                                bg='#243554', fg=self.fg_color, 
                                selectbackground=self.highlight_color,
                                highlightthickness=0, relief=tk.FLAT)
        self.slist.pack(fill=tk.BOTH, expand=True)
        self.scnt = self.create_styled_label(sf, "Selected: 0", 9, True)
        self.scnt.pack(anchor=tk.W, pady=5)
        
        bot = tk.Frame(self.tab1, bg=self.accent_color, height=60)
        bot.pack(fill=tk.X, padx=10, pady=10)
        bot.pack_propagate(False)
        
        self.status = tk.Label(bot, text="No photos loaded", font=('Arial', 10),
                               bg=self.accent_color, fg=self.fg_color)
        self.status.pack(side=tk.LEFT, padx=20, pady=15)
        
        self.create_styled_button(bot, "Generate Report", self.gen_report, 14).pack(side=tk.RIGHT, padx=10, pady=10)
        self.create_styled_button(bot, "Rename & Save", self.start_rename, 14).pack(side=tk.RIGHT, padx=5, pady=10)
        self.create_styled_button(bot, "Save Selected", self.save_selected, 14).pack(side=tk.RIGHT, padx=5, pady=10)


    def build_report_preview_tab(self):
        main = tk.PanedWindow(self.tab2, orient=tk.HORIZONTAL, bg=self.bg_color, sashwidth=5, sashrelief=tk.RAISED)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        left = self.create_styled_labelframe(main, "📋 Images for Report (sorted)")
        
        list_canvas = tk.Canvas(left, width=280, bg='#243554', highlightthickness=0)
        list_scrollbar = tk.Scrollbar(left, orient="vertical", command=list_canvas.yview,
                                       bg=self.accent_color, troughcolor='#243554')
        self.file_list_frame = tk.Frame(list_canvas, bg='#243554')
        
        list_canvas.configure(yscrollcommand=list_scrollbar.set)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        list_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_canvas.create_window((0,0), window=self.file_list_frame, anchor=tk.NW)
        
        self.file_list_frame.bind("<Configure>", lambda e: list_canvas.configure(scrollregion=list_canvas.bbox("all")))
        self.list_canvas = list_canvas
        
        right = tk.Frame(main, bg=self.bg_color)
        
        main.add(left, minsize=200)
        main.add(right, minsize=600)
        self.root.update_idletasks()
        main.sash_place(0, 280, 0)
        
        template_frame = self.create_styled_labelframe(right, "📂 Templates")
        template_frame.pack(fill=tk.X, pady=(0, 5))
        tf_top = tk.Frame(template_frame, bg=self.bg_color)
        tf_top.pack(fill=tk.X)
        self.create_styled_button(tf_top, "Upload Templates Folder", self.upload_templates, 22).pack(side=tk.LEFT, padx=5)
        self.template_status = tk.Label(tf_top, text="No templates loaded", font=('Arial', 9), bg=self.bg_color, fg='#ff9800')
        self.template_status.pack(side=tk.LEFT, padx=10)
        
        preview_frame = self.create_styled_labelframe(right, "📄 Report Page Preview")
        preview_frame.pack(fill=tk.BOTH, expand=True)
        
        nav = tk.Frame(preview_frame, bg=self.bg_color)
        nav.pack(fill=tk.X, pady=5)
        self.create_styled_button(nav, "◀◀ Prev", self.prev_page, 10).pack(side=tk.LEFT, padx=5)
        self.page_label = self.create_styled_label(nav, "Page 0 of 0", 12, True)
        self.page_label.pack(side=tk.LEFT, padx=20)
        self.create_styled_button(nav, "Next ▶▶", self.next_page, 10).pack(side=tk.LEFT, padx=5)
        tk.Label(nav, text="  Template:", font=('Arial', 10, 'bold'), bg=self.bg_color, fg=self.fg_color).pack(side=tk.LEFT, padx=(20, 5))
        self.template_var = tk.StringVar(value="1")
        self.template_combo = ttk.Combobox(nav, textvariable=self.template_var, width=25, state="readonly",
                                           values=["1: FRONT-WV + BACK-WV", "2: 1 WV + 2 CU", "3: 1 WV + 1 CU", "4: 4 CU"])
        self.template_combo.pack(side=tk.LEFT, padx=5)
        self.template_combo.bind("<<ComboboxSelected>>", self.on_template_select)
        self.create_styled_button(nav, "🔄 Refresh", self.refresh_preview, 10).pack(side=tk.RIGHT, padx=10)
        
        self.report_canvas = tk.Canvas(preview_frame, bg='#f5f5f5', highlightthickness=2,
                                        highlightbackground=self.accent_color)
        self.report_canvas.pack(fill=tk.BOTH, expand=True, pady=5)
        self.report_canvas.bind("<Configure>", lambda e: self.update_page_display())
        
        bot = tk.Frame(preview_frame, bg=self.bg_color)
        bot.pack(fill=tk.X, pady=5)
        self.create_styled_button(bot, "🖨️ Generate Multi-Page Report", self.gen_report_from_preview, 28).pack(side=tk.RIGHT, padx=5)

    def on_tab_change(self, event):
        if self.notebook.index(self.notebook.select()) == 1:
            self.refresh_preview()

    def upload_templates(self):
        folder = filedialog.askdirectory(title="Select Templates Folder")
        if not folder:
            return
        self.templates_folder.set(folder)
        self.templates = {}
        for f in os.listdir(folder):
            if f.lower().endswith('.docx'):
                fl = f.lower()
                if 'template_no_01' in fl or 'template_01' in fl:
                    self.templates[1] = os.path.join(folder, f)
                elif 'template_no_02' in fl or 'template_02' in fl:
                    self.templates[2] = os.path.join(folder, f)
                elif 'template_no_03' in fl or 'template_03' in fl:
                    self.templates[3] = os.path.join(folder, f)
                elif 'template_no_04' in fl or 'template_04' in fl:
                    self.templates[4] = os.path.join(folder, f)
        found = len(self.templates)
        if found == 4:
            self.template_status.config(text="All 4 templates loaded", fg='#4CAF50')
        elif found > 0:
            missing = [str(i) for i in range(1,5) if i not in self.templates]
            self.template_status.config(text=f"Found {found}/4. Missing: {', '.join(missing)}", fg='#ff9800')
        else:
            self.template_status.config(text="No templates found", fg='#f44336')

    def on_template_select(self, event=None):
        if not self.pages:
            return
        selection = self.template_var.get()
        template_num = int(selection.split(":")[0])
        self.page_templates[self.current_page] = template_num
        self.update_page_display()

    def get_page_template(self, page_idx):
        if page_idx in self.page_templates:
            return self.page_templates[page_idx]
        if page_idx < len(self.pages):
            page = self.pages[page_idx]
            images = page.get('images', [])
            if images:
                return self.detect_template(images)
        return 1

    def parse_filename(self, filepath):
        fname = os.path.basename(filepath).upper()
        run_match = re.search(r'RUN-?(\d+)', fname, re.IGNORECASE)
        run_num = int(run_match.group(1)) if run_match else 99
        if 'FRONT' in fname:
            side = 'FRONT'
            side_order = 0
        elif 'BACK' in fname:
            side = 'BACK'
            side_order = 1
        else:
            side = 'UNKNOWN'
            side_order = 9
        if '-WV' in fname or '_WV' in fname:
            view = 'WV'
            view_order = 0
        elif '-CU' in fname or '_CU' in fname:
            view = 'CU'
            view_order = 1
        else:
            view = 'UNKNOWN'
            view_order = 9
        return {
            'run': run_num,
            'side': side,
            'view': view,
            'side_order': side_order,
            'view_order': view_order,
            'path': filepath
        }

    def sort_images(self, image_list):
        def sort_key(fp):
            info = self.parse_filename(fp)
            return (info['run'], info['side_order'], info['view_order'])
        return sorted(image_list, key=sort_key)

    def detect_template(self, images):
        if not images:
            return 1
        infos = [self.parse_filename(fp) for fp in images]
        count = len(infos)
        sides = set(i['side'] for i in infos)
        views = [i['view'] for i in infos]
        wv_count = views.count('WV')
        cu_count = views.count('CU')
        
        if count == 2 and wv_count == 2:
            if 'FRONT' in sides and 'BACK' in sides:
                return 1
        if count == 3 and wv_count == 1 and cu_count == 2:
            if len(sides) == 1 and 'UNKNOWN' not in sides:
                return 2
        if count == 2 and wv_count == 1 and cu_count == 1:
            return 3
        if count == 4 and cu_count == 4:
            if len(sides) == 1 and 'UNKNOWN' not in sides:
                return 4
        return 1

    def group_images_by_run(self, image_list):
        groups = {}
        for fp in image_list:
            info = self.parse_filename(fp)
            run = info['run']
            if run not in groups:
                groups[run] = []
            groups[run].append(fp)
        return groups

    def refresh_preview(self):
        for w in self.file_list_frame.winfo_children():
            w.destroy()
        self.preview_checks = {}
        self.preview_thumbs = {}
        
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            lbl = tk.Label(self.file_list_frame, text="No photos selected.\nGo to Photo Selector tab\nand select photos.",
                          font=('Arial', 10), bg='#243554', fg='#aabbcc', justify=tk.CENTER)
            lbl.pack(pady=30)
            self.pages = []
            self.current_page = 0
            self.update_page_display()
            return
        
        self.sorted_images = self.sort_images(sel)
        
        current_run = None
        for i, fp in enumerate(self.sorted_images):
            info = self.parse_filename(fp)
            run_label = f"Run {str(info['run']).zfill(2)}"
            
            if run_label != current_run:
                current_run = run_label
                hdr = tk.Label(self.file_list_frame, text=f"─── {run_label} ───", 
                              font=('Arial', 9, 'bold'), bg='#243554', fg=self.highlight_color)
                hdr.pack(anchor=tk.W, pady=(15,5), padx=5)
            
            row = tk.Frame(self.file_list_frame, bg='#2a3f5f', padx=5, pady=3)
            row.pack(fill=tk.X, pady=2, padx=5)
            
            var = tk.BooleanVar(value=True)
            self.preview_checks[fp] = var
            chk = tk.Checkbutton(row, variable=var, command=self.update_pages,
                                bg='#2a3f5f', fg=self.fg_color, selectcolor='#1a2942',
                                activebackground='#2a3f5f', activeforeground=self.fg_color)
            chk.pack(side=tk.LEFT)
            
            category = f"{info['side']}-{info['view']}"
            cat_colors = {'FRONT-WV': '#4CAF50', 'BACK-WV': '#2196F3', 'FRONT-CU': '#FF9800', 'BACK-CU': '#E91E63', 'UNKNOWN-UNKNOWN': '#9E9E9E'}
            cat_lbl = tk.Label(row, text=f"[{category}]", font=('Arial', 8, 'bold'),
                              bg='#2a3f5f', fg=cat_colors.get(category, '#ffffff'), width=12)
            cat_lbl.pack(side=tk.LEFT)
            
            fname = os.path.basename(fp)
            name_lbl = tk.Label(row, text=fname[:20] + "..." if len(fname) > 20 else fname,
                               font=('Arial', 8), bg='#2a3f5f', fg=self.fg_color)
            name_lbl.pack(side=tk.LEFT, padx=5)
        
        self.update_pages()

    def update_pages(self):
        checked = [fp for fp in self.sorted_images if self.preview_checks.get(fp, tk.BooleanVar(value=False)).get()]
        
        self.pages = []
        run_groups = self.group_images_by_run(checked)
        
        for run_num in sorted(run_groups.keys()):
            run_images = run_groups[run_num]
            infos = [self.parse_filename(fp) for fp in run_images]
            
            wv_images = [i['path'] for i in infos if i['view'] == 'WV']
            cu_images = [i['path'] for i in infos if i['view'] == 'CU']
            
            page = {
                'images': run_images,
                'wv_images': wv_images,
                'cu_images': cu_images,
                'run': run_num
            }
            
            front_wv = [i['path'] for i in infos if i['side'] == 'FRONT' and i['view'] == 'WV']
            back_wv = [i['path'] for i in infos if i['side'] == 'BACK' and i['view'] == 'WV']
            page['front'] = front_wv[0] if front_wv else None
            page['back'] = back_wv[0] if back_wv else None
            page['closeups'] = cu_images
            
            self.pages.append(page)
        
        if self.current_page >= len(self.pages):
            self.current_page = max(0, len(self.pages) - 1)
        
        self.update_page_display()
# Chunk 3

    def update_page_display(self):
        self.report_canvas.delete("all")
        self.report_canvas.update_idletasks()
        
        # Get actual canvas size - dynamically fills available space
        cw = self.report_canvas.winfo_width()
        ch = self.report_canvas.winfo_height()
        if cw < 100: cw = 800
        if ch < 100: ch = 700
        
        margin = 15
        
        if not self.pages:
            self.page_label.config(text="Page 0 of 0")
            self.template_combo.set("1: FRONT-WV + BACK-WV")
            self.report_canvas.create_text(cw//2, ch//2, text="No pages to preview",
                                           font=('Arial', 18), fill='#666666')
            return
        
        self.page_label.config(text=f"Page {self.current_page + 1} of {len(self.pages)}")
        current_template = self.get_page_template(self.current_page)
        template_options = {1: "1: FRONT-WV + BACK-WV", 2: "2: 1 WV + 2 CU",
                           3: "3: 1 WV + 1 CU", 4: "4: 4 CU"}
        self.template_combo.set(template_options.get(current_template, template_options[1]))
        
        page = self.pages[self.current_page]
        
        # Shadow and main preview box - fills canvas dynamically
        self.report_canvas.create_rectangle(margin+5, margin+5, cw-margin, ch-margin, fill='#cccccc', outline='')
        self.report_canvas.create_rectangle(margin, margin, cw-margin-5, ch-margin-5, fill='white', outline='#1a2942', width=2)
        
        # Calculate usable area
        box_left = margin + 10
        box_right = cw - margin - 15
        box_top = margin + 15
        box_bottom = ch - margin - 20
        box_width = box_right - box_left
        box_height = box_bottom - box_top
        center_x = (box_left + box_right) // 2
        header_height = 28
        
        if current_template == 1:
            # Template 1: FRONT-WV on top, BACK-WV on bottom
            half_height = (box_height - header_height * 2 - 20) // 2
            
            # FRONT header
            self.report_canvas.create_rectangle(box_left, box_top, box_right, box_top + header_height, fill='#4CAF50', outline='')
            self.report_canvas.create_text(center_x, box_top + header_height//2, text="FRONT (Wide-View)", font=('Arial', 11, 'bold'), fill='white')
            
            # FRONT image area
            front_top = box_top + header_height + 5
            front_bottom = front_top + half_height
            self.report_canvas.create_rectangle(box_left, front_top, box_right, front_bottom, outline="#cccccc", dash=(3,3))
            if page.get('front'):
                self.draw_preview_image(page['front'], center_x, (front_top + front_bottom)//2, max_w=box_width-20, max_h=half_height-20)
            else:
                self.report_canvas.create_text(center_x, (front_top + front_bottom)//2, text="(No Front image)", font=('Arial', 13), fill='#cccccc')
            
            # BACK header
            back_header_top = front_bottom + 10
            self.report_canvas.create_rectangle(box_left, back_header_top, box_right, back_header_top + header_height, fill='#2196F3', outline='')
            self.report_canvas.create_text(center_x, back_header_top + header_height//2, text="BACK (Wide-View)", font=('Arial', 11, 'bold'), fill='white')
            
            # BACK image area
            back_top = back_header_top + header_height + 5
            back_bottom = box_bottom
            self.report_canvas.create_rectangle(box_left, back_top, box_right, back_bottom, outline="#cccccc", dash=(3,3))
            if page.get('back'):
                self.draw_preview_image(page['back'], center_x, (back_top + back_bottom)//2, max_w=box_width-20, max_h=back_bottom-back_top-20)
            else:
                self.report_canvas.create_text(center_x, (back_top + back_bottom)//2, text="(No Back image)", font=('Arial', 13), fill='#cccccc')
        
        elif current_template == 3:
            # Template 3: 1 WV on top, 1 CU on bottom
            half_height = (box_height - header_height * 2 - 20) // 2
            
            # WIDE-VIEW header
            self.report_canvas.create_rectangle(box_left, box_top, box_right, box_top + header_height, fill='#4CAF50', outline='')
            self.report_canvas.create_text(center_x, box_top + header_height//2, text="WIDE-VIEW (Row 2)", font=('Arial', 11, 'bold'), fill='white')
            
            # WIDE-VIEW image area
            wv_top = box_top + header_height + 5
            wv_bottom = wv_top + half_height
            self.report_canvas.create_rectangle(box_left, wv_top, box_right, wv_bottom, outline="#cccccc", dash=(3,3))
            wv_imgs = page.get('wv_images', [])
            if wv_imgs:
                self.draw_preview_image(wv_imgs[0], center_x, (wv_top + wv_bottom)//2, max_w=box_width-20, max_h=half_height-20)
            else:
                self.report_canvas.create_text(center_x, (wv_top + wv_bottom)//2, text="(No Wide-view)", font=('Arial', 13), fill='#cccccc')
            
            # CLOSE-UP header
            cu_header_top = wv_bottom + 10
            self.report_canvas.create_rectangle(box_left, cu_header_top, box_right, cu_header_top + header_height, fill='#FF9800', outline='')
            self.report_canvas.create_text(center_x, cu_header_top + header_height//2, text="CLOSE-UP (Row 4)", font=('Arial', 11, 'bold'), fill='white')
            
            # CLOSE-UP image area
            cu_top = cu_header_top + header_height + 5
            cu_bottom = box_bottom
            self.report_canvas.create_rectangle(box_left, cu_top, box_right, cu_bottom, outline="#cccccc", dash=(3,3))
            closeups = page.get('closeups', [])
            if closeups:
                self.draw_preview_image(closeups[0], center_x, (cu_top + cu_bottom)//2, max_w=box_width-20, max_h=cu_bottom-cu_top-20)
            else:
                self.report_canvas.create_text(center_x, (cu_top + cu_bottom)//2, text="(No Close-up)", font=('Arial', 13), fill='#cccccc')
        
        elif current_template == 2:
            # Template 2: 1 WV on top, 2 CU side by side on bottom
            top_height = int((box_height - header_height * 2 - 20) * 0.45)
            bottom_height = box_height - header_height * 2 - 20 - top_height
            
            # WIDE-VIEW header
            self.report_canvas.create_rectangle(box_left, box_top, box_right, box_top + header_height, fill='#4CAF50', outline='')
            self.report_canvas.create_text(center_x, box_top + header_height//2, text="WIDE-VIEW (Row 2)", font=('Arial', 11, 'bold'), fill='white')
            
            # WIDE-VIEW image area
            wv_top = box_top + header_height + 5
            wv_bottom = wv_top + top_height
            self.report_canvas.create_rectangle(box_left, wv_top, box_right, wv_bottom, outline="#cccccc", dash=(3,3))
            wv_imgs = page.get('wv_images', [])
            if wv_imgs:
                self.draw_preview_image(wv_imgs[0], center_x, (wv_top + wv_bottom)//2, max_w=box_width-20, max_h=top_height-20)
            else:
                self.report_canvas.create_text(center_x, (wv_top + wv_bottom)//2, text="(No Wide-view)", font=('Arial', 13), fill='#cccccc')
            
            # CLOSE-UPS header
            cu_header_top = wv_bottom + 10
            self.report_canvas.create_rectangle(box_left, cu_header_top, box_right, cu_header_top + header_height, fill='#FF9800', outline='')
            self.report_canvas.create_text(center_x, cu_header_top + header_height//2, text="CLOSE-UPS (Row 4 - 2 cells)", font=('Arial', 11, 'bold'), fill='white')
            
            # CLOSE-UP image areas (2 side by side)
            cu_top = cu_header_top + header_height + 5
            cu_bottom = box_bottom
            half_width = (box_width - 10) // 2
            
            # Left close-up
            self.report_canvas.create_rectangle(box_left, cu_top, box_left + half_width, cu_bottom, outline="#cccccc", dash=(3,3))
            # Right close-up
            self.report_canvas.create_rectangle(box_left + half_width + 10, cu_top, box_right, cu_bottom, outline="#cccccc", dash=(3,3))
            
            closeups = page.get('closeups', [])
            left_center = box_left + half_width // 2
            right_center = box_left + half_width + 10 + half_width // 2
            cu_center_y = (cu_top + cu_bottom) // 2
            
            if len(closeups) >= 1:
                self.draw_preview_image(closeups[0], left_center, cu_center_y, max_w=half_width-20, max_h=cu_bottom-cu_top-20)
            else:
                self.report_canvas.create_text(left_center, cu_center_y, text="(Close-up 1)", font=('Arial', 11), fill='#cccccc')
            if len(closeups) >= 2:
                self.draw_preview_image(closeups[1], right_center, cu_center_y, max_w=half_width-20, max_h=cu_bottom-cu_top-20)
            else:
                self.report_canvas.create_text(right_center, cu_center_y, text="(Close-up 2)", font=('Arial', 11), fill='#cccccc')
        
        elif current_template == 4:
            # Template 4: 4 CU in 2x2 grid
            half_height = (box_height - header_height * 2 - 20) // 2
            half_width = (box_width - 10) // 2
            
            # Top row header
            self.report_canvas.create_rectangle(box_left, box_top, box_right, box_top + header_height, fill='#FF9800', outline='')
            self.report_canvas.create_text(center_x, box_top + header_height//2, text="CLOSE-UPS (Row 2 - 2 cells)", font=('Arial', 11, 'bold'), fill='white')
            
            # Top row image areas
            top_img_top = box_top + header_height + 5
            top_img_bottom = top_img_top + half_height - header_height
            self.report_canvas.create_rectangle(box_left, top_img_top, box_left + half_width, top_img_bottom, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(box_left + half_width + 10, top_img_top, box_right, top_img_bottom, outline="#cccccc", dash=(3,3))
            
            # Bottom row header
            bottom_header_top = top_img_bottom + 10
            self.report_canvas.create_rectangle(box_left, bottom_header_top, box_right, bottom_header_top + header_height, fill='#FF9800', outline='')
            self.report_canvas.create_text(center_x, bottom_header_top + header_height//2, text="CLOSE-UPS (Row 4 - 2 cells)", font=('Arial', 11, 'bold'), fill='white')
            
            # Bottom row image areas
            bottom_img_top = bottom_header_top + header_height + 5
            bottom_img_bottom = box_bottom
            self.report_canvas.create_rectangle(box_left, bottom_img_top, box_left + half_width, bottom_img_bottom, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(box_left + half_width + 10, bottom_img_top, box_right, bottom_img_bottom, outline="#cccccc", dash=(3,3))
            
            closeups = page.get('closeups', [])
            left_center = box_left + half_width // 2
            right_center = box_left + half_width + 10 + half_width // 2
            top_center_y = (top_img_top + top_img_bottom) // 2
            bottom_center_y = (bottom_img_top + bottom_img_bottom) // 2
            img_max_w = half_width - 20
            img_max_h_top = top_img_bottom - top_img_top - 20
            img_max_h_bottom = bottom_img_bottom - bottom_img_top - 20
            
            positions = [(left_center, top_center_y, img_max_h_top), (right_center, top_center_y, img_max_h_top),
                        (left_center, bottom_center_y, img_max_h_bottom), (right_center, bottom_center_y, img_max_h_bottom)]
            
            for idx, (cx, cy, max_h) in enumerate(positions):
                if idx < len(closeups):
                    self.draw_preview_image(closeups[idx], cx, cy, max_w=img_max_w, max_h=max_h)
                else:
                    self.report_canvas.create_text(cx, cy, text=f"(Close-up {idx+1})", font=('Arial', 11), fill='#cccccc')

    def draw_preview_image(self, filepath, cx, cy, max_w, max_h):
        try:
            img = Image.open(filepath)
            img.thumbnail((max_w, max_h))
            tkim = ImageTk.PhotoImage(img)
            self.preview_thumbs[filepath + "_preview"] = tkim
            self.report_canvas.create_image(cx, cy, image=tkim)
        except:
            self.report_canvas.create_text(cx, cy, text="Error loading image", fill="red")

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.update_page_display()

    def next_page(self):
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            self.update_page_display()

    def replace_image(self, old_path):
        new_path = filedialog.askopenfilename(title="Select Replacement Image",
            filetypes=[("Images", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if not new_path: return
        idx = self.sorted_images.index(old_path)
        del self.preview_checks[old_path]
        self.sorted_images[idx] = new_path
        if old_path in self.checks:
            old_val = self.checks[old_path].get()
            del self.checks[old_path]
            self.checks[new_path] = tk.BooleanVar(value=old_val)
        self.refresh_preview()
        messagebox.showinfo("Replaced", f"Image replaced with:\n{os.path.basename(new_path)}")
# Chunk 4

    def gen_report_from_preview(self):
        if not self.pages:
            messagebox.showwarning("Warning", "No pages to generate.")
            return
        if not self.templates:
            messagebox.showwarning("Warning", "Please upload templates folder first.")
            return
        if not DOCX_OK:
            messagebox.showerror("Error", "python-docx not installed.")
            return
        needed = set(self.get_page_template(i) for i in range(len(self.pages)))
        missing = [str(t) for t in needed if t not in self.templates]
        if missing:
            messagebox.showerror("Error", f"Missing template(s): {', '.join(missing)}")
            return
        op = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word","*.docx")], title="Save Report As")
        if not op:
            return
        try:
            first_tpl = self.get_page_template(0)
            master = Document(self.templates[first_tpl])
            vid = self.vb_id.get().strip() or "Unknown"
            self.fill_template_page(master, self.pages[0], first_tpl, vid)
            for idx in range(1, len(self.pages)):
                tpl = self.get_page_template(idx)
                page_doc = Document(self.templates[tpl])
                self.fill_template_page(page_doc, self.pages[idx], tpl, vid)
                master.add_page_break()
                for elem in page_doc.element.body:
                    master.element.body.append(elem)
            master.save(op)
            messagebox.showinfo("Success", f"Report saved with {len(self.pages)} pages!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")

    def fill_template_page(self, doc, page, tpl_num, vb_id):
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if 'vb_id_txt' in cell.text:
                        cell.text = cell.text.replace('vb_id_txt', vb_id)
        def ins_img(cell, path, h=2.2):
            cell.paragraphs[0].clear()
            cell.paragraphs[0].add_run().add_picture(path, height=Inches(h))
        if tpl_num == 1 and len(doc.tables) >= 3:
            if page.get('front'):
                ins_img(doc.tables[2].rows[1].cells[0], page['front'])
            if page.get('back'):
                ins_img(doc.tables[2].rows[3].cells[0], page['back'])
        elif tpl_num == 2 and len(doc.tables) >= 3:
            wv_imgs = page.get('wv_images', [])
            if wv_imgs:
                ins_img(doc.tables[2].rows[1].cells[0], wv_imgs[0])
            closeups = page.get('closeups', [])
            if len(closeups) >= 1 and len(doc.tables[2].rows) >= 4:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[0], 1.8)
            if len(closeups) >= 2 and len(doc.tables[2].rows[3].cells) >= 2:
                ins_img(doc.tables[2].rows[3].cells[1], closeups[1], 1.8)
        elif tpl_num == 3 and len(doc.tables) >= 3:
            wv_imgs = page.get('wv_images', [])
            if wv_imgs:
                ins_img(doc.tables[2].rows[1].cells[0], wv_imgs[0])
            closeups = page.get('closeups', [])
            if closeups:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[0])
        elif tpl_num == 4 and len(doc.tables) >= 3:
            closeups = page.get('closeups', [])
            if len(closeups) >= 1:
                ins_img(doc.tables[2].rows[1].cells[0], closeups[0], 1.8)
            if len(closeups) >= 2 and len(doc.tables[2].rows[1].cells) >= 2:
                ins_img(doc.tables[2].rows[1].cells[1], closeups[1], 1.8)
            if len(closeups) >= 3 and len(doc.tables[2].rows) >= 4:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[2], 1.8)
            if len(closeups) >= 4 and len(doc.tables[2].rows[3].cells) >= 2:
                ins_img(doc.tables[2].rows[3].cells[1], closeups[3], 1.8)

    def browse(self):
        f = filedialog.askdirectory()
        if f: self.path.set(f)

    def sel_all(self):
        for v in self.checks.values(): v.set(True)
        self.upd_sel()

    def clr_all(self):
        for v in self.checks.values(): v.set(False)
        self.upd_sel()

    def load(self):
        p = self.path.get()
        if not p or not os.path.isdir(p): return
        for w in self.inner.winfo_children(): w.destroy()
        self.photos, self.thumbs, self.checks = [], {}, {}
        for f in sorted(os.listdir(p)):
            if f.lower().endswith(('.jpg','.jpeg','.png','.gif','.bmp')):
                self.photos.append(os.path.join(p, f))
        r, c = 0, 0
        for ph in self.photos:
            fr = tk.Frame(self.inner, bg='#243554', padx=5, pady=5)
            fr.grid(row=r, column=c, padx=5, pady=5)
            try:
                im = Image.open(ph)
                im.thumbnail((100,100))
                tkim = ImageTk.PhotoImage(im)
                self.thumbs[ph] = tkim
                btn = tk.Button(fr, image=tkim, command=lambda x=ph:self.show(x),
                               bg='#243554', relief=tk.FLAT, cursor='hand2')
                btn.pack()
            except: pass
            lbl = tk.Label(fr, text=os.path.basename(ph)[:12], font=('Arial', 8),
                          bg='#243554', fg=self.fg_color)
            lbl.pack()
            v = tk.BooleanVar()
            self.checks[ph] = v
            chk = tk.Checkbutton(fr, text="Select", variable=v, command=self.upd_sel,
                                bg='#243554', fg=self.fg_color, selectcolor='#1a2942',
                                activebackground='#243554', font=('Arial', 8))
            chk.pack()
            c += 1
            if c > 5: c, r = 0, r+1
        self.status.config(text=f"Loaded {len(self.photos)} photos")

    def show(self, path):
        if self.crop_mode: self.cancel_crop()
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.current_photo_path = path
        try:
            self.original_image = Image.open(path)
            w, h = self.original_image.size
            self.lbl_n.config(text="Name: "+os.path.basename(path))
            self.lbl_s.config(text=f"Size: {w}x{h}")
            self.upd_pv()
        except: pass

    def upd_pv(self):
        if not self.original_image: return
        im = self.original_image.copy()
        im.thumbnail((int(500*self.zoom), int(500*self.zoom)))
        self.preview_img = ImageTk.PhotoImage(im)
        self.pcv.delete("all")
        self.pcv.create_image(260+self.pan_x, 175+self.pan_y, image=self.preview_img)

    def upd_sel(self):
        self.slist.delete(0, tk.END)
        sel = [p for p,v in self.checks.items() if v.get()]
        self.scnt.config(text=f"Selected: {len(sel)}")
        for p in sel: self.slist.insert(tk.END, os.path.basename(p))

    def zoom_in(self):
        self.zoom = min(5, self.zoom*1.3)
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def zoom_out(self):
        self.zoom = max(0.3, self.zoom/1.3)
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def zoom_reset(self):
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.crop_coords = [50,50,200,200]
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def pan(self, dx, dy):
        self.pan_x += dx
        self.pan_y += dy
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def toggle_crop(self):
        if not self.original_image:
            messagebox.showwarning("Warning", "Load a photo first")
            return
        self.crop_mode = not self.crop_mode
        if self.crop_mode:
            self.crop_btn.config(text="Cropping...")
            self.crop_frame.pack(fill=tk.X, pady=3)
            self.crop_coords = [50,50,200,200]
            self.upd_pv()
            self.draw_crop()
            self.pcv.bind("<Button-1>", self.crop_down)
            self.pcv.bind("<B1-Motion>", self.crop_drag)
            self.pcv.bind("<ButtonRelease-1>", self.crop_up)
        else:
            self.cancel_crop()

    def draw_crop(self):
        self.pcv.delete("crop")
        x1,y1,x2,y2 = self.crop_coords
        self.pcv.create_rectangle(x1,y1,x2,y2, outline="red", width=2, tags="crop", dash=(4,2))
        for i,(hx,hy) in enumerate([(x1,y1),(x2,y1),(x1,y2),(x2,y2)]):
            self.pcv.create_rectangle(hx-6,hy-6,hx+6,hy+6, fill="#ff6b6b", outline="white", tags="crop")

    def crop_down(self, e):
        x,y = e.x, e.y
        x1,y1,x2,y2 = self.crop_coords
        for i,(hx,hy) in enumerate([(x1,y1),(x2,y1),(x1,y2),(x2,y2)]):
            if abs(x-hx)<12 and abs(y-hy)<12:
                self.active_handle, self.crop_start = i, (x,y)
                return
        if min(x1,x2)<x<max(x1,x2) and min(y1,y2)<y<max(y1,y2):
            self.active_handle, self.crop_start = "move", (x,y)

    def crop_drag(self, e):
        if self.crop_start is None: return
        dx, dy = e.x-self.crop_start[0], e.y-self.crop_start[1]
        x1,y1,x2,y2 = self.crop_coords
        if self.active_handle=="move": self.crop_coords = [x1+dx,y1+dy,x2+dx,y2+dy]
        elif self.active_handle==0: self.crop_coords = [x1+dx,y1+dy,x2,y2]
        elif self.active_handle==1: self.crop_coords = [x1,y1+dy,x2+dx,y2]
        elif self.active_handle==2: self.crop_coords = [x1+dx,y1,x2,y2+dy]
        elif self.active_handle==3: self.crop_coords = [x1,y1,x2+dx,y2+dy]
        self.crop_start = (e.x, e.y)
        self.draw_crop()

    def crop_up(self, e):
        self.active_handle, self.crop_start = None, None
# Chunk 5

    def apply_crop(self):
        if not self.original_image: return
        x1,y1,x2,y2 = self.crop_coords
        cx, cy = 260+self.pan_x, 175+self.pan_y
        im = self.original_image.copy()
        im.thumbnail((int(500*self.zoom), int(500*self.zoom)))
        pw, ph = im.size
        ox, oy = cx-pw//2, cy-ph//2
        rx1, ry1 = max(0,x1-ox), max(0,y1-oy)
        rx2, ry2 = min(pw,x2-ox), min(ph,y2-oy)
        ow, oh = self.original_image.size
        sx, sy = ow/pw, oh/ph
        fx1, fy1 = int(rx1*sx), int(ry1*sy)
        fx2, fy2 = int(rx2*sx), int(ry2*sy)
        if fx2<=fx1 or fy2<=fy1:
            messagebox.showwarning("Warning","Invalid crop region")
            return
        self.original_image = self.original_image.crop((fx1,fy1,fx2,fy2))
        self.cancel_crop()
        self.upd_pv()
        w,h = self.original_image.size
        self.lbl_s.config(text=f"Size: {w}x{h} (cropped)")

    def cancel_crop(self):
        self.crop_mode = False
        self.crop_btn.config(text="Crop")
        self.crop_frame.pack_forget()
        self.pcv.delete("crop")
        self.pcv.unbind("<Button-1>")
        self.pcv.unbind("<B1-Motion>")
        self.pcv.unbind("<ButtonRelease-1>")

    def save_selected(self):
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showwarning("Warning", "No photos selected")
            return
        out = filedialog.askdirectory(title="Select Output Folder")
        if not out: return
        self.output_folder = out
        for p in sel:
            shutil.copy2(p, out)
        messagebox.showinfo("Saved", f"Copied {len(sel)} photos to {out}")

    def start_rename(self):
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showwarning("Warning", "No photos selected")
            return
        out = filedialog.askdirectory(title="Select Output Folder")
        if not out: return
        self.output_folder = out
        self.renamed_photos = []
        for p in sel:
            self.rename_one(p)
        if self.renamed_photos:
            messagebox.showinfo("Done", f"Renamed {len(self.renamed_photos)} photos")

    def rename_one(self, path):
        win = tk.Toplevel(self.root)
        win.title("Rename Photo")
        win.geometry("500x400")
        win.configure(bg=self.bg_color)
        win.transient(self.root)
        win.grab_set()
        
        try:
            im = Image.open(path)
            im.thumbnail((300,250))
            tkim = ImageTk.PhotoImage(im)
            lbl = tk.Label(win, image=tkim, bg=self.bg_color)
            lbl.image = tkim
            lbl.pack(pady=10)
        except: pass
        
        tk.Label(win, text=f"Current: {os.path.basename(path)}", font=('Arial', 9),
                bg=self.bg_color, fg=self.fg_color).pack()
        
        tk.Label(win, text="New name:", font=('Arial', 10, 'bold'),
                bg=self.bg_color, fg=self.fg_color).pack(pady=(10,5))
        
        base, ext = os.path.splitext(os.path.basename(path))
        prefix = self.vb_id.get().strip() + "_" if self.vb_id.get().strip() else ""
        nv = tk.StringVar(value=prefix+base)
        ent = self.create_styled_entry(win, nv, 40)
        ent.pack(pady=5)
        
        result = {'saved': False}
        
        def save():
            nn = nv.get().strip()
            if nn:
                np_path = os.path.join(self.output_folder, nn+ext)
                shutil.copy2(path, np_path)
                self.renamed_photos.append(np_path)
                result['saved'] = True
            win.destroy()
        
        def skip():
            win.destroy()
        
        bf = tk.Frame(win, bg=self.bg_color)
        bf.pack(pady=15)
        self.create_styled_button(bf, "Save", save, 10).pack(side=tk.LEFT, padx=10)
        self.create_styled_button(bf, "Skip", skip, 10).pack(side=tk.LEFT, padx=10)
        
        win.wait_window()

    def gen_report(self):
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showwarning("Warning", "No photos selected")
            return
        messagebox.showinfo("Info", "Use Report Preview tab for multi-page reports with templates.")

if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()

    