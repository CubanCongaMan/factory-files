#!/usr/bin/env python3
# 03/30/2026 Photo Selector Version 25
#Author: RLS & Factory AI Droids
#Fix: Updated deprecated trace() to trace_add() for Tkinter compatibility
#Fix: Widened preview panel to 40% using PanedWindow for resizable layout

import os, shutil, re, tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk

try:
    import ttkbootstrap as ttk
    from ttkbootstrap.constants import *
    THEME_AVAILABLE = True
except ImportError:
    from tkinter import ttk
    THEME_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docxcompose.composer import Composer
    DOCX_OK = True
except:
    DOCX_OK = False

class App:
    def __init__(self, root):
        self.root = root
        root.title("Photo Selector v25")
        root.geometry("1400x850")
        
        self.bg_color = "#1a2942"
        self.fg_color = "#ffffff"
        self.accent_color = "#3d5a80"
        self.highlight_color = "#4a90d9"
        
        self.photos, self.thumbs, self.checks = [], {}, {}
        self.preview_img, self.original_image, self.current_photo_path = None, None, None
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.output_folder, self.renamed_photos = "", []
        self.crop_mode, self.crop_coords, self.crop_start, self.active_handle = False, [50,50,200,200], None, None
        self.vb_id = tk.StringVar()
        self.path = tk.StringVar()
        
        self.sorted_images = []
        self.preview_checks = {}
        self.preview_thumbs = {}
        self.current_page = 0
        self.pages = []
        
        self.templates_folder = tk.StringVar()
        self.templates = {}
        self.page_templates = {}
        
        self.style = ttk.Style()
        if not THEME_AVAILABLE:
            self.configure_fallback_style()
        
        main_container = tk.Frame(root, bg=self.bg_color)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        header = tk.Frame(main_container, bg=self.accent_color, height=50)
        header.pack(fill=tk.X, padx=0, pady=0)
        header.pack_propagate(False)
        
        title_label = tk.Label(header, text="📷 PHOTO SELECTOR & REPORT GENERATOR", 
                               font=('Arial', 16, 'bold'), bg=self.accent_color, fg=self.fg_color)
        title_label.pack(side=tk.LEFT, padx=20, pady=10)
        
        version_label = tk.Label(header, text="v25", font=('Arial', 10), 
                                  bg=self.accent_color, fg="#aabbcc")
        version_label.pack(side=tk.RIGHT, padx=20, pady=15)
        
        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.tab1 = tk.Frame(self.notebook, bg=self.bg_color)
        self.notebook.add(self.tab1, text="  📁 Photo Selector  ")
        self.build_photo_selector_tab()
        
        self.tab2 = tk.Frame(self.notebook, bg=self.bg_color)
        self.notebook.add(self.tab2, text="  📄 Report Preview  ")
        self.build_report_preview_tab()
        
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_change)

    def configure_fallback_style(self):
        self.style.configure('TNotebook', background=self.bg_color)
        self.style.configure('TNotebook.Tab', padding=[20, 10], font=('Arial', 10, 'bold'))
        self.style.map('TNotebook.Tab', 
                       background=[('selected', self.highlight_color), ('!selected', self.accent_color)],
                       foreground=[('selected', self.fg_color), ('!selected', '#cccccc')])

    def create_styled_label(self, parent, text, font_size=10, bold=False):
        font_weight = 'bold' if bold else 'normal'
        return tk.Label(parent, text=text, font=('Arial', font_size, font_weight),
                       bg=self.bg_color, fg=self.fg_color)

    def create_styled_button(self, parent, text, command, width=12):
        btn = tk.Button(parent, text=text, command=command, width=width,
                       font=('Arial', 9, 'bold'), bg=self.accent_color, fg=self.fg_color,
                       activebackground=self.highlight_color, activeforeground=self.fg_color,
                       relief=tk.FLAT, cursor='hand2', padx=10, pady=5)
        btn.bind('<Enter>', lambda e: btn.config(bg=self.highlight_color))
        btn.bind('<Leave>', lambda e: btn.config(bg=self.accent_color))
        return btn

    def create_styled_entry(self, parent, textvariable, width=40):
        entry = tk.Entry(parent, textvariable=textvariable, width=width,
                        font=('Arial', 10), bg='#2a3f5f', fg=self.fg_color,
                        insertbackground=self.fg_color, relief=tk.FLAT,
                        highlightthickness=2, highlightcolor=self.highlight_color,
                        highlightbackground=self.accent_color)
        return entry

    def create_styled_labelframe(self, parent, text):
        frame = tk.LabelFrame(parent, text=text, font=('Arial', 10, 'bold'),
                             bg=self.bg_color, fg=self.fg_color, padx=10, pady=10,
                             relief=tk.GROOVE, bd=2)
        return frame

    def build_photo_selector_tab(self):
        top = tk.Frame(self.tab1, bg=self.bg_color, pady=10)
        top.pack(fill=tk.X, padx=10)
        
        self.create_styled_label(top, "Boring ID:", 10, True).pack(side=tk.LEFT)
        self.create_styled_entry(top, self.vb_id, 15).pack(side=tk.LEFT, padx=(5, 20))
        
        self.create_styled_label(top, "Folder:", 10, True).pack(side=tk.LEFT)
        self.create_styled_entry(top, self.path, 35).pack(side=tk.LEFT, padx=5)
        
        self.create_styled_button(top, "Browse", self.browse, 8).pack(side=tk.LEFT, padx=3)
        self.create_styled_button(top, "Load", self.load, 8).pack(side=tk.LEFT, padx=3)
        self.create_styled_button(top, "Select All", self.sel_all, 10).pack(side=tk.LEFT, padx=(15,3))
        self.create_styled_button(top, "Clear All", self.clr_all, 10).pack(side=tk.LEFT, padx=3)
        
        main = tk.PanedWindow(self.tab1, orient=tk.HORIZONTAL, bg=self.bg_color, sashwidth=5, sashrelief=tk.RAISED)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        left = self.create_styled_labelframe(main, "📷 Photos - Click to Preview")
        
        self.canvas = tk.Canvas(left, bg='#243554', highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        self.inner = tk.Frame(self.canvas, bg='#243554')
        self.canvas.create_window((0,0), window=self.inner, anchor=tk.NW)
        
        right = tk.Frame(main, bg=self.bg_color)
        
        main.add(left, minsize=300)
        main.add(right, minsize=400)
        self.root.update_idletasks()
        main.sash_place(0, 800, 0)
        
        pf = self.create_styled_labelframe(right, "🔍 Preview")
        pf.pack(fill=tk.BOTH, expand=True)
        
        self.pcv = tk.Canvas(pf, bg='#243554', width=520, height=350, highlightthickness=0)
        self.pcv.pack(fill=tk.BOTH, expand=True)
        
        zf = tk.Frame(pf, bg=self.bg_color)
        zf.pack(fill=tk.X, pady=5)
        self.create_styled_button(zf, "Zoom+", self.zoom_in, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(zf, "Zoom-", self.zoom_out, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(zf, "Reset", self.zoom_reset, 7).pack(side=tk.LEFT, padx=2)
        self.crop_btn = self.create_styled_button(zf, "Crop", self.toggle_crop, 7)
        self.crop_btn.pack(side=tk.LEFT, padx=10)
        
        self.crop_frame = tk.Frame(pf, bg=self.bg_color)
        self.create_styled_button(self.crop_frame, "Apply", self.apply_crop, 7).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(self.crop_frame, "Cancel", self.cancel_crop, 7).pack(side=tk.LEFT, padx=2)
        
        pnf = tk.Frame(pf, bg=self.bg_color)
        pnf.pack(fill=tk.X, pady=5)
        self.create_styled_button(pnf, "▲", lambda:self.pan(0,-30), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "▼", lambda:self.pan(0,30), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "◀", lambda:self.pan(-30,0), 4).pack(side=tk.LEFT, padx=2)
        self.create_styled_button(pnf, "▶", lambda:self.pan(30,0), 4).pack(side=tk.LEFT, padx=2)
        
        self.lbl_n = self.create_styled_label(pf, "Name: -", 9)
        self.lbl_n.pack(anchor=tk.W, pady=2)
        self.lbl_s = self.create_styled_label(pf, "Size: -", 9)
        self.lbl_s.pack(anchor=tk.W, pady=2)
        
        sf = self.create_styled_labelframe(right, "✓ Selected Photos")
        sf.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        self.slist = tk.Listbox(sf, height=5, font=('Arial', 9),
                                bg='#243554', fg=self.fg_color, 
                                selectbackground=self.highlight_color,
                                highlightthickness=0, relief=tk.FLAT)
        self.slist.pack(fill=tk.BOTH, expand=True)
        self.scnt = self.create_styled_label(sf, "Selected: 0", 9, True)
        self.scnt.pack(anchor=tk.W, pady=5)
        
        bot = tk.Frame(self.tab1, bg=self.accent_color, height=60)
        bot.pack(fill=tk.X, padx=10, pady=10)
        bot.pack_propagate(False)
        
        self.status = tk.Label(bot, text="No photos loaded", font=('Arial', 10),
                               bg=self.accent_color, fg=self.fg_color)
        self.status.pack(side=tk.LEFT, padx=20, pady=15)
        
        self.create_styled_button(bot, "Generate Report", self.gen_report, 14).pack(side=tk.RIGHT, padx=10, pady=10)
        self.create_styled_button(bot, "Rename & Save", self.start_rename, 14).pack(side=tk.RIGHT, padx=5, pady=10)
        self.create_styled_button(bot, "Save Selected", self.save_selected, 14).pack(side=tk.RIGHT, padx=5, pady=10)

       def build_report_preview_tab(self):
        main = tk.PanedWindow(self.tab2, orient=tk.HORIZONTAL, bg=self.bg_color, sashwidth=5, sashrelief=tk.RAISED)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        left = self.create_styled_labelframe(main, "📋 Images for Report (sorted)")
        
        list_canvas = tk.Canvas(left, width=320, bg='#243554', highlightthickness=0)
        list_scrollbar = tk.Scrollbar(left, orient="vertical", command=list_canvas.yview,
                                       bg=self.accent_color, troughcolor='#243554')
        self.file_list_frame = tk.Frame(list_canvas, bg='#243554')
        
        list_canvas.configure(yscrollcommand=list_scrollbar.set)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        list_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_canvas.create_window((0,0), window=self.file_list_frame, anchor=tk.NW)
        
        self.file_list_frame.bind("<Configure>", lambda e: list_canvas.configure(scrollregion=list_canvas.bbox("all")))
        self.list_canvas = list_canvas
        
        right = tk.Frame(main, bg=self.bg_color)
        
        main.add(left, minsize=250)
        main.add(right, minsize=300)
        self.root.update_idletasks()
        main.sash_place(0, 350, 0)
        
        template_frame = self.create_styled_labelframe(right, "📂 Templates")

        template_frame.pack(fill=tk.X, pady=(0, 10))
        tf_top = tk.Frame(template_frame, bg=self.bg_color)
        tf_top.pack(fill=tk.X)
        self.create_styled_button(tf_top, "Upload Templates Folder", self.upload_templates, 22).pack(side=tk.LEFT, padx=5)
        self.template_status = tk.Label(tf_top, text="No templates loaded", font=('Arial', 9), bg=self.bg_color, fg='#ff9800')
        self.template_status.pack(side=tk.LEFT, padx=10)
        
        preview_frame = self.create_styled_labelframe(right, "📄 Report Page Preview")
        preview_frame.pack(fill=tk.BOTH, expand=True)
        
        nav = tk.Frame(preview_frame, bg=self.bg_color)
        nav.pack(fill=tk.X, pady=10)
        self.create_styled_button(nav, "◀◀ Prev", self.prev_page, 10).pack(side=tk.LEFT, padx=5)
        self.page_label = self.create_styled_label(nav, "Page 0 of 0", 12, True)
        self.page_label.pack(side=tk.LEFT, padx=20)
        self.create_styled_button(nav, "Next ▶▶", self.next_page, 10).pack(side=tk.LEFT, padx=5)
        tk.Label(nav, text="  Template:", font=('Arial', 10, 'bold'), bg=self.bg_color, fg=self.fg_color).pack(side=tk.LEFT, padx=(20, 5))
        self.template_var = tk.StringVar(value="1")
        self.template_combo = ttk.Combobox(nav, textvariable=self.template_var, width=25, state="readonly",
                                           values=["1: 2 Wide-view (Front+Back)", "2: 1 Wide + 2 Close-ups", "3: 1 Wide + 1 Close-up", "4: 4 Close-ups"])
        self.template_combo.pack(side=tk.LEFT, padx=5)
        self.template_combo.bind("<<ComboboxSelected>>", self.on_template_select)
        self.create_styled_button(nav, "🔄 Refresh", self.refresh_preview, 10).pack(side=tk.RIGHT, padx=10)
        
        self.report_canvas = tk.Canvas(preview_frame, bg='#f5f5f5', width=350, height=450, highlightthickness=2,
                                        highlightbackground=self.accent_color)
        self.report_canvas.pack(fill=tk.BOTH, expand=True, pady
        
        bot = tk.Frame(preview_frame, bg=self.bg_color)
        bot.pack(fill=tk.X, pady=10)
        self.create_styled_button(bot, "🖨️ Generate Multi-Page Report", self.gen_report_from_preview, 28).pack(side=tk.RIGHT, padx=5)

    def on_tab_change(self, event):
        if self.notebook.index(self.notebook.select()) == 1:
            self.refresh_preview()

    def upload_templates(self):
        folder = filedialog.askdirectory(title="Select Templates Folder")
        if not folder:
            return
        self.templates_folder.set(folder)
        self.templates = {}
        for f in os.listdir(folder):
            if f.lower().endswith('.docx'):
                fl = f.lower()
             
                if 'template_no_01' in fl or 'template_01' in fl:
                    self.templates[1] = os.path.join(folder, f)
                elif 'template_no_02' in fl or 'template_02' in fl:
                    self.templates[2] = os.path.join(folder, f)
                elif 'template_no_03' in fl or 'template_03' in fl:
                    self.templates[3] = os.path.join(folder, f)
                elif 'template_no_04' in fl or 'template_04' in fl:
                    self.templates[4] = os.path.join(folder, f)
        found = len(self.templates)
        if found == 4:
            self.template_status.config(text="All 4 templates loaded", fg='#4CAF50')
        elif found > 0:
            missing = [str(i) for i in range(1,5) if i not in self.templates]
            self.template_status.config(text=f"Found {found}/4. Missing: {', '.join(missing)}", fg='#ff9800')
        else:
            self.template_status.config(text="No templates found", fg='#f44336')

    def on_template_select(self, event=None):
        if not self.pages:
            return
        selection = self.template_var.get()
        template_num = int(selection.split(":")[0])
        self.page_templates[self.current_page] = template_num
        self.update_page_display()

    def get_page_template(self, page_idx):
        if page_idx in self.page_templates:
            return self.page_templates[page_idx]
        if page_idx < len(self.pages):
            page = self.pages[page_idx]
            if page.get('front') and page.get('back'):
                return 1
        return 1

    def parse_filename(self, filepath):
        fname = os.path.basename(filepath).lower()
        run_match = re.search(r'_run(\d+)(-?[abc])?_', fname, re.IGNORECASE)
        run_num = int(run_match.group(1)) if run_match else 99
        suffix = run_match.group(2).replace('-','').upper() if run_match and run_match.group(2) else ''
        suffix_order = {'': 0, 'A': 1, 'B': 2, 'C': 3}
        suffix_val = suffix_order.get(suffix, 0)
        if '_front' in fname:
            cat_order = 0
            category = 'Front'
        elif '_back' in fname:
            cat_order = 1
            category = 'Back'
        elif '_close' in fname:
            cat_order = 2
            category = 'Close-up'
        else:
            cat_order = 9
            category = 'Unknown'
        return (run_num, suffix_val, cat_order, category, suffix)

    def sort_images(self, image_list):
        def sort_key(fp):
            run_num, suffix_val, cat_order, _, _ = self.parse_filename(fp)
            return (run_num, suffix_val, cat_order)
        return sorted(image_list, key=sort_key)

    def refresh_preview(self):
        for w in self.file_list_frame.winfo_children():
            w.destroy()
        self.preview_checks = {}
        self.preview_thumbs = {}
        
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            lbl = tk.Label(self.file_list_frame, text="No photos selected.\nGo to Photo Selector tab\nand select photos.",
                          font=('Arial', 10), bg='#243554', fg='#aabbcc', justify=tk.CENTER)
            lbl.pack(pady=30)
            self.pages = []
            self.current_page = 0
            self.update_page_display()
            return
        
        self.sorted_images = self.sort_images(sel)
        
        current_run = None
        for i, fp in enumerate(self.sorted_images):
            run_num, suffix_val, cat_order, category, suffix = self.parse_filename(fp)
            run_label = f"Run {str(run_num).zfill(2)}" + (f"-{suffix}" if suffix else "")
            
            if run_label != current_run:
                current_run = run_label
                hdr = tk.Label(self.file_list_frame, text=f"─── {run_label} ───", 
                              font=('Arial', 9, 'bold'), bg='#243554', fg=self.highlight_color)
                hdr.pack(anchor=tk.W, pady=(15,5), padx=5)
            
            row = tk.Frame(self.file_list_frame, bg='#2a3f5f', padx=5, pady=3)
            row.pack(fill=tk.X, pady=2, padx=5)
            
            var = tk.BooleanVar(value=True)
            self.preview_checks[fp] = var
            chk = tk.Checkbutton(row, variable=var, command=self.update_pages,
                                bg='#2a3f5f', fg=self.fg_color, selectcolor='#1a2942',
                                activebackground='#2a3f5f', activeforeground=self.fg_color)
            chk.pack(side=tk.LEFT)
            
            cat_colors = {'Front': '#4CAF50', 'Back': '#2196F3', 'Close-up': '#FF9800', 'Unknown': '#9E9E9E'}
            cat_lbl = tk.Label(row, text=f"[{category}]", font=('Arial', 8, 'bold'),
                              bg='#2a3f5f', fg=cat_colors.get(category, '#ffffff'), width=10)
            cat_lbl.pack(side=tk.LEFT)
            
            fname = os.path.basename(fp)
            name_lbl = tk.Label(row, text=fname[:22] + "..." if len(fname) > 22 else fname,
                               font=('Arial', 8), bg='#2a3f5f', fg=self.fg_color)
            name_lbl.pack(side=tk.LEFT, padx=5)
            
            rep_btn = tk.Button(row, text="Replace", font=('Arial', 7), width=6,
                               bg=self.accent_color, fg=self.fg_color, relief=tk.FLAT,
                               command=lambda f=fp: self.replace_image(f))
            rep_btn.pack(side=tk.RIGHT, padx=2)
        
        self.update_pages()

    def update_pages(self):
        checked = [fp for fp in self.sorted_images if self.preview_checks.get(fp, tk.BooleanVar(value=False)).get()]
        
        self.pages = []
        i = 0
        while i < len(checked):
            fp = checked[i]
            _, _, cat_order, category, _ = self.parse_filename(fp)
            
            if category == 'Front':
                page = {'front': fp, 'back': None, 'closeups': []}
                if i + 1 < len(checked):
                    next_fp = checked[i + 1]
                    _, _, next_cat, next_category, _ = self.parse_filename(next_fp)
                    if next_category == 'Back':
                        page['back'] = next_fp
                        i += 1
                self.pages.append(page)
            elif category == 'Back':
                self.pages.append({'front': None, 'back': fp, 'closeups': []})
            elif category == 'Close-up':
                self.pages.append({'front': None, 'back': None, 'closeups': [fp]})
            else:
                self.pages.append({'unknown': fp})
            i += 1
        
        if self.current_page >= len(self.pages):
            self.current_page = max(0, len(self.pages) - 1)
        
        self.update_page_display()

    def update_page_display(self):
        self.report_canvas.delete("all")
        
        if not self.pages:
            self.page_label.config(text="Page 0 of 0")
            self.template_combo.set("1: 2 Wide-view (Front+Back)")
            self.report_canvas.create_text(350, 225, text="No pages to preview", 
                                           font=('Arial', 16), fill='#666666')
            return
        
        self.page_label.config(text=f"Page {self.current_page + 1} of {len(self.pages)}")
        current_template = self.get_page_template(self.current_page)
        template_options = {1: "1: 2 Wide-view (Front+Back)", 2: "2: 1 Wide + 2 Close-ups",
                           3: "3: 1 Wide + 1 Close-up", 4: "4: 4 Close-ups"}
        self.template_combo.set(template_options.get(current_template, template_options[1]))
        
        page = self.pages[self.current_page]
        self.report_canvas.create_rectangle(55, 25, 655, 435, fill='#cccccc', outline='')
        self.report_canvas.create_rectangle(50, 20, 650, 430, fill='white', outline='#1a2942', width=2)
        
        if current_template == 1:
            self.report_canvas.create_rectangle(60, 35, 640, 55, fill='#4CAF50', outline='')
            self.report_canvas.create_text(350, 45, text="FRONT (Wide-view)", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 58, 640, 215, outline="#cccccc", dash=(3,3))
            if page.get('front'):
                self.draw_preview_image(page['front'], 350, 135, max_w=560, max_h=145)
                self.report_canvas.create_text(350, 205, text=os.path.basename(page['front']), font=('Arial', 8, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(350, 135, text="(No Front image)", font=('Arial', 11), fill='#cccccc')
            self.report_canvas.create_rectangle(60, 225, 640, 245, fill='#2196F3', outline='')
            self.report_canvas.create_text(350, 235, text="BACK (Wide-view)", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 248, 640, 420, outline="#cccccc", dash=(3,3))
            if page.get('back'):
                self.draw_preview_image(page['back'], 350, 332, max_w=560, max_h=160)
                self.report_canvas.create_text(350, 410, text=os.path.basename(page['back']), font=('Arial', 8, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(350, 332, text="(No Back image)", font=('Arial', 11), fill='#cccccc')
        elif current_template == 3:
            self.report_canvas.create_rectangle(60, 35, 640, 55, fill='#4CAF50', outline='')
            self.report_canvas.create_text(350, 45, text="WIDE-VIEW", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 58, 640, 215, outline="#cccccc", dash=(3,3))
            wide_img = page.get('front') or page.get('back')
            if wide_img:
                self.draw_preview_image(wide_img, 350, 135, max_w=560, max_h=145)
                self.report_canvas.create_text(350, 205, text=os.path.basename(wide_img), font=('Arial', 8, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(350, 135, text="(No Wide-view)", font=('Arial', 11), fill='#cccccc')
            self.report_canvas.create_rectangle(60, 225, 640, 245, fill='#FF9800', outline='')
            self.report_canvas.create_text(350, 235, text="CLOSE-UP", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 248, 640, 420, outline="#cccccc", dash=(3,3))
            closeups = page.get('closeups', [])
            if closeups:
                self.draw_preview_image(closeups[0], 350, 332, max_w=560, max_h=160)
                self.report_canvas.create_text(350, 410, text=os.path.basename(closeups[0]), font=('Arial', 8, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(350, 332, text="(No Close-up)", font=('Arial', 11), fill='#cccccc')
        elif current_template == 2:
            self.report_canvas.create_rectangle(60, 35, 640, 55, fill='#4CAF50', outline='')
            self.report_canvas.create_text(350, 45, text="WIDE-VIEW", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 58, 640, 200, outline="#cccccc", dash=(3,3))
            wide_img = page.get('front') or page.get('back')
            if wide_img:
                self.draw_preview_image(wide_img, 350, 125, max_w=560, max_h=130)
                self.report_canvas.create_text(350, 190, text=os.path.basename(wide_img), font=('Arial', 8, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(350, 125, text="(No Wide-view)", font=('Arial', 11), fill='#cccccc')
            self.report_canvas.create_rectangle(60, 210, 640, 230, fill='#FF9800', outline='')
            self.report_canvas.create_text(350, 220, text="CLOSE-UPS (2 images)", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 235, 345, 420, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(355, 235, 640, 420, outline="#cccccc", dash=(3,3))
            closeups = page.get('closeups', [])
            if len(closeups) >= 1:
                self.draw_preview_image(closeups[0], 202, 320, max_w=270, max_h=170)
                self.report_canvas.create_text(202, 410, text=os.path.basename(closeups[0])[:20], font=('Arial', 7, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(202, 320, text="(Close-up 1)", font=('Arial', 10), fill='#cccccc')
            if len(closeups) >= 2:
                self.draw_preview_image(closeups[1], 498, 320, max_w=270, max_h=170)
                self.report_canvas.create_text(498, 410, text=os.path.basename(closeups[1])[:20], font=('Arial', 7, 'bold'), fill='#000000')
            else:
                self.report_canvas.create_text(498, 320, text="(Close-up 2)", font=('Arial', 10), fill='#cccccc')
        elif current_template == 4:
            self.report_canvas.create_rectangle(60, 35, 640, 55, fill='#FF9800', outline='')
            self.report_canvas.create_text(350, 45, text="CLOSE-UPS (Row 1)", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 58, 345, 210, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(355, 58, 640, 210, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(60, 220, 640, 240, fill='#FF9800', outline='')
            self.report_canvas.create_text(350, 230, text="CLOSE-UPS (Row 2)", font=('Arial', 10, 'bold'), fill='white')
            self.report_canvas.create_rectangle(60, 243, 345, 420, outline="#cccccc", dash=(3,3))
            self.report_canvas.create_rectangle(355, 243, 640, 420, outline="#cccccc", dash=(3,3))
            closeups = page.get('closeups', [])
            positions = [(202, 130), (498, 130), (202, 328), (498, 328)]
            labels_y = [200, 200, 410, 410]
            for idx, (cx, cy) in enumerate(positions):
                if idx < len(closeups):
                    self.draw_preview_image(closeups[idx], cx, cy, max_w=270, max_h=140)
                    self.report_canvas.create_text(cx, labels_y[idx], text=os.path.basename(closeups[idx])[:20], font=('Arial', 7, 'bold'), fill='#000000')
                else:
                    self.report_canvas.create_text(cx, cy, text=f"(Close-up {idx+1})", font=('Arial', 10), fill='#cccccc')

    def draw_preview_image(self, filepath, cx, cy, max_w, max_h):
        try:
            img = Image.open(filepath)
            img.thumbnail((max_w, max_h))
            tkim = ImageTk.PhotoImage(img)
            self.preview_thumbs[filepath + "_preview"] = tkim
            self.report_canvas.create_image(cx, cy, image=tkim)
        except:
            self.report_canvas.create_text(cx, cy, text="Error loading image", fill="red")

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.update_page_display()

    def next_page(self):
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            self.update_page_display()

    def replace_image(self, old_path):
        new_path = filedialog.askopenfilename(title="Select Replacement Image",
            filetypes=[("Images", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if not new_path: return
        new_fname = os.path.basename(new_path).lower()
        if '_front' not in new_fname and '_back' not in new_fname and '_close' not in new_fname:
            messagebox.showwarning("Warning", "Filename must contain _front, _back, or _close.")
            return
        idx = self.sorted_images.index(old_path)
        del self.preview_checks[old_path]
        self.sorted_images[idx] = new_path
        if old_path in self.checks:
            old_val = self.checks[old_path].get()
            del self.checks[old_path]
            self.checks[new_path] = tk.BooleanVar(value=old_val)
        self.refresh_preview()
        messagebox.showinfo("Replaced", f"Image replaced with:\n{os.path.basename(new_path)}")

    def gen_report_from_preview(self):
        if not self.pages:
            messagebox.showwarning("Warning", "No pages to generate.")
            return
        if not self.templates:
            messagebox.showwarning("Warning", "Please upload templates folder first.")
            return
        if not DOCX_OK:
            messagebox.showerror("Error", "python-docx not installed.")
            return
        needed = set(self.get_page_template(i) for i in range(len(self.pages)))
        missing = [str(t) for t in needed if t not in self.templates]
        if missing:
            messagebox.showerror("Error", f"Missing template(s): {', '.join(missing)}")
            return
        op = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word","*.docx")], title="Save Report As")
        if not op:
            return
        try:
            first_tpl = self.get_page_template(0)
            master = Document(self.templates[first_tpl])
            vid = self.vb_id.get().strip() or "Unknown"
            self.fill_template_page(master, self.pages[0], first_tpl, vid)
            for idx in range(1, len(self.pages)):
                tpl = self.get_page_template(idx)
                page_doc = Document(self.templates[tpl])
                self.fill_template_page(page_doc, self.pages[idx], tpl, vid)
                master.add_page_break()
                for elem in page_doc.element.body:
                    master.element.body.append(elem)
            master.save(op)
            messagebox.showinfo("Success", f"Report saved with {len(self.pages)} pages!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")

    def fill_template_page(self, doc, page, tpl_num, vb_id):
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if 'vb_id_txt' in cell.text:
                        cell.text = cell.text.replace('vb_id_txt', vb_id)
        def ins_img(cell, path, h=2.2):
            cell.paragraphs[0].clear()
            cell.paragraphs[0].add_run().add_picture(path, height=Inches(h))
        if tpl_num == 1 and len(doc.tables) >= 3:
            if page.get('front'):
                ins_img(doc.tables[2].rows[1].cells[0], page['front'])
            if page.get('back'):
                ins_img(doc.tables[2].rows[3].cells[0], page['back'])
        elif tpl_num == 2 and len(doc.tables) >= 3:
            wide = page.get('front') or page.get('back')
            if wide:
                ins_img(doc.tables[2].rows[1].cells[0], wide)
            closeups = page.get('closeups', [])
            if len(closeups) >= 1 and len(doc.tables[2].rows) >= 4:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[0], 1.8)
            if len(closeups) >= 2 and len(doc.tables[2].rows[3].cells) >= 2:
                ins_img(doc.tables[2].rows[3].cells[1], closeups[1], 1.8)
        elif tpl_num == 3 and len(doc.tables) >= 3:
            wide = page.get('front') or page.get('back')
            if wide:
                ins_img(doc.tables[2].rows[1].cells[0], wide)
            closeups = page.get('closeups', [])
            if closeups:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[0])
        elif tpl_num == 4 and len(doc.tables) >= 3:
            closeups = page.get('closeups', [])
            if len(closeups) >= 1:
                ins_img(doc.tables[2].rows[1].cells[0], closeups[0], 1.8)
            if len(closeups) >= 2 and len(doc.tables[2].rows[1].cells) >= 2:
                ins_img(doc.tables[2].rows[1].cells[1], closeups[1], 1.8)
            if len(closeups) >= 3 and len(doc.tables[2].rows) >= 4:
                ins_img(doc.tables[2].rows[3].cells[0], closeups[2], 1.8)
            if len(closeups) >= 4 and len(doc.tables[2].rows[3].cells) >= 2:
                ins_img(doc.tables[2].rows[3].cells[1], closeups[3], 1.8)

    def browse(self):
        f = filedialog.askdirectory()
        if f: self.path.set(f)

    def sel_all(self):
        for v in self.checks.values(): v.set(True)
        self.upd_sel()

    def clr_all(self):
        for v in self.checks.values(): v.set(False)
        self.upd_sel()

    def load(self):
        p = self.path.get()
        if not p or not os.path.isdir(p): return
        for w in self.inner.winfo_children(): w.destroy()
        self.photos, self.thumbs, self.checks = [], {}, {}
        for f in sorted(os.listdir(p)):
            if f.lower().endswith(('.jpg','.jpeg','.png','.gif','.bmp')):
                self.photos.append(os.path.join(p, f))
        r, c = 0, 0
        for ph in self.photos:
            fr = tk.Frame(self.inner, bg='#243554', padx=5, pady=5)
            fr.grid(row=r, column=c, padx=5, pady=5)
            try:
                im = Image.open(ph)
                im.thumbnail((100,100))
                tkim = ImageTk.PhotoImage(im)
                self.thumbs[ph] = tkim
                btn = tk.Button(fr, image=tkim, command=lambda x=ph:self.show(x),
                               bg='#243554', relief=tk.FLAT, cursor='hand2')
                btn.pack()
            except: pass
            lbl = tk.Label(fr, text=os.path.basename(ph)[:12], font=('Arial', 8),
                          bg='#243554', fg=self.fg_color)
            lbl.pack()
            v = tk.BooleanVar()
            self.checks[ph] = v
            chk = tk.Checkbutton(fr, text="Select", variable=v, command=self.upd_sel,
                                bg='#243554', fg=self.fg_color, selectcolor='#1a2942',
                                activebackground='#243554', font=('Arial', 8))
            chk.pack()
            c += 1
            if c > 5: c, r = 0, r+1
        self.status.config(text=f"Loaded {len(self.photos)} photos")

    def show(self, path):
        if self.crop_mode: self.cancel_crop()
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.current_photo_path = path
        try:
            self.original_image = Image.open(path)
            w, h = self.original_image.size
            self.lbl_n.config(text="Name: "+os.path.basename(path))
            self.lbl_s.config(text=f"Size: {w}x{h}")
            self.upd_pv()
        except: pass

    def upd_pv(self):
        if not self.original_image: return
        im = self.original_image.copy()
        im.thumbnail((int(500*self.zoom), int(500*self.zoom)))
        self.preview_img = ImageTk.PhotoImage(im)
        self.pcv.delete("all")
        self.pcv.create_image(260+self.pan_x, 175+self.pan_y, image=self.preview_img)

    def upd_sel(self):
        self.slist.delete(0, tk.END)
        sel = [p for p,v in self.checks.items() if v.get()]
        self.scnt.config(text=f"Selected: {len(sel)}")
        for p in sel: self.slist.insert(tk.END, os.path.basename(p))

    def zoom_in(self):
        self.zoom = min(5, self.zoom*1.3)
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def zoom_out(self):
        self.zoom = max(0.3, self.zoom/1.3)
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def zoom_reset(self):
        self.zoom, self.pan_x, self.pan_y = 1.0, 0, 0
        self.crop_coords = [50,50,200,200]
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def pan(self, dx, dy):
        self.pan_x += dx
        self.pan_y += dy
        self.upd_pv()
        if self.crop_mode: self.draw_crop()

    def toggle_crop(self):
        if not self.original_image:
            messagebox.showwarning("Warning", "Load a photo first")
            return
        self.crop_mode = not self.crop_mode
        if self.crop_mode:
            self.crop_btn.config(text="Cropping...")
            self.crop_frame.pack(fill=tk.X, pady=3)
            self.crop_coords = [50,50,200,200]
            self.upd_pv()
            self.draw_crop()
            self.pcv.bind("<Button-1>", self.crop_down)
            self.pcv.bind("<B1-Motion>", self.crop_drag)
            self.pcv.bind("<ButtonRelease-1>", self.crop_up)
        else:
            self.cancel_crop()

    def draw_crop(self):
        self.pcv.delete("crop")
        x1,y1,x2,y2 = self.crop_coords
        self.pcv.create_rectangle(x1,y1,x2,y2, outline="red", width=2, tags="crop", dash=(4,2))
        for i,(hx,hy) in enumerate([(x1,y1),(x2,y1),(x1,y2),(x2,y2)]):
            self.pcv.create_rectangle(hx-6,hy-6,hx+6,hy+6, fill="#ff6b6b", outline="white", tags="crop")

    def crop_down(self, e):
        x,y = e.x, e.y
        x1,y1,x2,y2 = self.crop_coords
        for i,(hx,hy) in enumerate([(x1,y1),(x2,y1),(x1,y2),(x2,y2)]):
            if abs(x-hx)<12 and abs(y-hy)<12:
                self.active_handle, self.crop_start = i, (x,y)
                return
        if min(x1,x2)<x<max(x1,x2) and min(y1,y2)<y<max(y1,y2):
            self.active_handle, self.crop_start = "move", (x,y)

    def crop_drag(self, e):
        if self.crop_start is None: return
        dx, dy = e.x-self.crop_start[0], e.y-self.crop_start[1]
        x1,y1,x2,y2 = self.crop_coords
        if self.active_handle=="move": self.crop_coords = [x1+dx,y1+dy,x2+dx,y2+dy]
        elif self.active_handle==0: self.crop_coords = [x1+dx,y1+dy,x2,y2]
        elif self.active_handle==1: self.crop_coords = [x1,y1+dy,x2+dx,y2]
        elif self.active_handle==2: self.crop_coords = [x1+dx,y1,x2,y2+dy]
        elif self.active_handle==3: self.crop_coords = [x1,y1,x2+dx,y2+dy]
        self.crop_start = (e.x, e.y)
        self.draw_crop()

    def crop_up(self, e):
        self.active_handle, self.crop_start = None, None

    def apply_crop(self):
        if not self.original_image: return
        iw, ih = self.original_image.size
        md = int(500*self.zoom)
        sc = min(md/iw, md/ih)
        dw, dh = int(iw*sc), int(ih*sc)
        il, it = 260+self.pan_x-dw//2, 175+self.pan_y-dh//2
        x1,y1,x2,y2 = self.crop_coords
        if x1>x2: x1,x2 = x2,x1
        if y1>y2: y1,y2 = y2,y1
        rx1,ry1 = max(0,min(1,(x1-il)/dw)), max(0,min(1,(y1-it)/dh))
        rx2,ry2 = max(0,min(1,(x2-il)/dw)), max(0,min(1,(y2-it)/dh))
        px1,py1,px2,py2 = int(rx1*iw),int(ry1*ih),int(rx2*iw),int(ry2*ih)
        if px2<=px1 or py2<=py1:
            messagebox.showwarning("Warning","Invalid crop")
            return
        cropped = self.original_image.crop((px1,py1,px2,py2))
        ch = messagebox.askyesnocancel("Save","Yes=Replace\nNo=Save new\nCancel=Abort")
        if ch is None: return
        if ch: cropped.save(self.current_photo_path)
        else:
            sp = filedialog.asksaveasfilename(defaultextension=".jpg")
            if sp: cropped.save(sp)
        self.cancel_crop()
        self.original_image = Image.open(self.current_photo_path)
        self.upd_pv()

    def cancel_crop(self):
        self.crop_mode = False
        self.crop_btn.config(text="Crop")
        self.crop_frame.pack_forget()
        self.pcv.delete("crop")
        self.pcv.unbind("<Button-1>")
        self.pcv.unbind("<B1-Motion>")
        self.pcv.unbind("<ButtonRelease-1>")
        self.upd_pv()

    def save_selected(self):
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showinfo("Info","No photos selected")
            return
        self.output_folder = filedialog.askdirectory(title="Output Folder")
        if not self.output_folder: return
        saved, non_compliant = 0, []
        for fp in sel:
            fname = os.path.basename(fp)
            dest = os.path.join(self.output_folder, fname)
            shutil.copy2(fp, dest)
            saved += 1
            fl = fname.lower()
            if '_front' not in fl and '_back' not in fl and '_close' not in fl:
                non_compliant.append(fname)
        msg = f"Saved {saved} files to:\n{self.output_folder}"
        if non_compliant:
            msg += f"\n\nWARNING: {len(non_compliant)} file(s) don't match naming convention"
        messagebox.showinfo("Save Complete", msg)

    def start_rename(self):
        if not self.vb_id.get().strip():
            messagebox.showwarning("Warning","Enter Boring ID")
            return
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showinfo("Info","No photos selected")
            return
        self.output_folder = filedialog.askdirectory(title="Output Folder")
        if not self.output_folder: return
        self.renamed_photos = []
        for i, ph in enumerate(sel):
            self.rename_one(ph, i+1, len(sel))

    def rename_one(self, ph, num, total):
        d = tk.Toplevel(self.root)
        d.title(f"Rename {num}/{total}")
        d.geometry("800x650")
        d.configure(bg=self.bg_color)
        d.grab_set()
        tk.Label(d, text="Original: "+os.path.basename(ph), font=('Arial', 10),
                bg=self.bg_color, fg=self.fg_color).pack(pady=10)
        cv = tk.Canvas(d, bg='#243554', width=400, height=400, highlightthickness=0)
        cv.pack(pady=10)
        try:
            im = Image.open(ph)
            im.thumbnail((380,380))
            d.tki = ImageTk.PhotoImage(im)
            cv.create_image(200,200,image=d.tki)
        except: pass
        of = tk.Frame(d, bg=self.bg_color, pady=10)
        of.pack(fill=tk.X, padx=20)
        tk.Label(of, text="Run#:", font=('Arial', 10, 'bold'),
                bg=self.bg_color, fg=self.fg_color).pack(side=tk.LEFT)
        rn = tk.StringVar(value="01")
        rc = ttk.Combobox(of, textvariable=rn, width=4, state="readonly", 
                         values=[str(i).zfill(2) for i in range(1,21)])
        rc.pack(side=tk.LEFT, padx=5)
        tk.Label(of, text="Suffix:", font=('Arial', 10, 'bold'),
                bg=self.bg_color, fg=self.fg_color).pack(side=tk.LEFT, padx=(15,0))
        rs = tk.StringVar(value="")
        ttk.Combobox(of, textvariable=rs, width=3, state="readonly", 
                    values=["","A","B","C"]).pack(side=tk.LEFT, padx=5)
        tk.Label(of, text="Category:", font=('Arial', 10, 'bold'),
                bg=self.bg_color, fg=self.fg_color).pack(side=tk.LEFT, padx=(15,0))
        ct = tk.StringVar(value="Front")
        for txt in ["Front", "Back", "Close-up"]:
            tk.Radiobutton(of, text=txt, variable=ct, value=txt,
                          bg=self.bg_color, fg=self.fg_color, selectcolor='#1a2942',
                          activebackground=self.bg_color, font=('Arial', 9)).pack(side=tk.LEFT)
        nm = tk.StringVar()
        tk.Label(d, textvariable=nm, font=('Arial', 11, 'bold'),
                bg=self.bg_color, fg='#4CAF50').pack(pady=10)
        def upn(*a):
            suf = "-"+rs.get() if rs.get() else ""
            nm.set(f"{self.vb_id.get()}_Run{rn.get()}{suf}_{ct.get()}{os.path.splitext(ph)[1].lower()}")
        rn.trace_add("write", upn)
        rs.trace_add("write", upn)
        ct.trace_add("write", upn)
        upn()
        bf = tk.Frame(d, bg=self.bg_color)
        bf.pack(pady=15)
        def sav():
            np = os.path.join(self.output_folder, nm.get())
            shutil.copy2(ph, np)
            self.renamed_photos.append(np)
            d.destroy()
        self.create_styled_button(bf, "Save", sav, 10).pack(side=tk.LEFT, padx=10)
        self.create_styled_button(bf, "Skip", d.destroy, 10).pack(side=tk.LEFT, padx=10)
        d.wait_window()

    def gen_report(self):
        if not DOCX_OK:
            messagebox.showerror("Error","Install python-docx")
            return
        tp = filedialog.askopenfilename(title="Select Template", filetypes=[("Word","*.docx")])
        if not tp: return
        sel = [p for p,v in self.checks.items() if v.get()]
        if not sel:
            messagebox.showwarning("Warning","No photos selected.")
            return
        front, back, closeup = [], [], []
        for fp in sel:
            fl = os.path.basename(fp).lower()
            if '_front' in fl: front.append(fp)
            elif '_back' in fl: back.append(fp)
            elif '_close' in fl: closeup.append(fp)
        if not any([front,back,closeup]):
            messagebox.showwarning("Warning","Selected photos must have _front, _back, or _close in filename")
            return
        op = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word","*.docx")])
        if not op: return
        try:
            doc = Document(tp)
            vid = self.vb_id.get().strip() or "Unknown"
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if 'vb_id_txt' in cell.text: cell.text = cell.text.replace('vb_id_txt',vid)
            def insert_img(cell, img_path):
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, height=Inches(2.2))
            def set_label(cell, label, fname):
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"{label}: {fname}")
                r.font.size = Pt(8)
            if len(doc.tables) >= 3:
                if front:
                    set_label(doc.tables[2].rows[0].cells[0], "Front", os.path.basename(front[0]))
                    insert_img(doc.tables[2].rows[1].cells[0], front[0])
                if back:
                    set_label(doc.tables[2].rows[2].cells[0], "Back", os.path.basename(back[0]))
                    insert_img(doc.tables[2].rows[3].cells[0], back[0])
            if len(doc.tables) >= 6 and closeup:
                set_label(doc.tables[5].rows[0].cells[0], "Closeup", os.path.basename(closeup[0]))
                insert_img(doc.tables[5].rows[1].cells[0], closeup[0])
            doc.save(op)
            messagebox.showinfo("Done",f"Report saved!\nFront: {len(front)}\nBack: {len(back)}\nCloseup: {len(closeup)}")
        except Exception as e:
            messagebox.showerror("Error",str(e))

if __name__ == "__main__":
    if THEME_AVAILABLE:
        root = ttk.Window(themename="superhero")
    else:
        root = tk.Tk()
    App(root)
    root.mainloop()

        
        
