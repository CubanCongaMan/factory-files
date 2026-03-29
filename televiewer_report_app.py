#!/usr/bin/env python3
"""
================================================================================
    OPTICAL TELEVIEWER PHOTO REPORT GENERATOR
    Version 1.0.0
================================================================================

    Generate professional photo reports from downhole televiewer images.
    
    For each borehole:
    1. Upload a televiewer image
    2. Enter header information (or pull from database)
    3. Generate a complete Word report with cropped depth images

    Copyright (c) 2026 Star Engineering, Inc.
    Developed with Factory AI

================================================================================
"""

import os
import sys
from datetime import datetime
from pathlib import Path

# Check dependencies
def check_dependencies():
    """Check and report missing dependencies."""
    missing = []
    
    try:
        from PIL import Image
    except ImportError:
        missing.append("Pillow")
    
    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")
    
    try:
        import requests
    except ImportError:
        missing.append("requests")
    
    if missing:
        print("\n" + "="*60)
        print("  MISSING DEPENDENCIES")
        print("="*60)
        print(f"\n  Please install the following packages:\n")
        print(f"    pip install {' '.join(missing)}")
        print("\n" + "="*60)
        sys.exit(1)

check_dependencies()

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """Application configuration."""
    
    # App Info
    APP_NAME = "Optical Televiewer Photo Report Generator"
    VERSION = "1.0.0"
    COMPANY = "Star Engineering, Inc."
    
    # Supabase
    SUPABASE_URL = "https://oebbsdcdnnzqwdoacyyz.supabase.co"
    
    # Image Processing (calibrated values)
    PIXELS_PER_FOOT = 360
    DEPTH_OFFSET = -74
    HEADER_HEIGHT = 110
    RIGHT_CROP_X = 685
    LABEL_OFFSET = 20
    IMAGE_HEIGHT_INCHES = 7.0
    
    # Report Notes
    NOTES = """Notes:
(1) Camera type and recording settings are as follows:
Type: Digital Optical Borehole Televiewer (Mount Sopris)
Frame Rate: 20 HZ
Exposure: Medium
Lighting: 20
Resolution: 900
Logging Speed: 2.7-2.9 ft/min
(2) Photograph Not-to-Scale
(3) Borehole orientation presented in the above image is relative to the local magnetic field"""


# =============================================================================
# CONSOLE INTERFACE
# =============================================================================

class Console:
    """Console UI helper."""
    
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BOLD = '\033[1m'
    END = '\033[0m'
    
    @classmethod
    def clear(cls):
        os.system('cls' if os.name == 'nt' else 'clear')
    
    @classmethod
    def header(cls):
        cls.clear()
        print(f"""
{cls.CYAN}{'='*65}{cls.END}
{cls.BOLD}{cls.CYAN}  OPTICAL TELEVIEWER PHOTO REPORT GENERATOR{cls.END}
{cls.CYAN}  Version {Config.VERSION} | {Config.COMPANY}{cls.END}
{cls.CYAN}{'='*65}{cls.END}
""")
    
    @classmethod
    def section(cls, title):
        print(f"\n{cls.BOLD}{cls.CYAN}  {title}{cls.END}")
        print(f"  {'-'*50}")
    
    @classmethod
    def success(cls, msg):
        print(f"  {cls.GREEN}✓{cls.END} {msg}")
    
    @classmethod
    def error(cls, msg):
        print(f"  {cls.RED}✗{cls.END} {msg}")
    
    @classmethod
    def info(cls, msg):
        print(f"  {cls.CYAN}→{cls.END} {msg}")
    
    @classmethod
    def warning(cls, msg):
        print(f"  {cls.YELLOW}!{cls.END} {msg}")
    
    @classmethod
    def prompt(cls, msg, default=None):
        if default:
            result = input(f"  {cls.YELLOW}>{cls.END} {msg} [{default}]: ").strip()
            return result if result else default
        return input(f"  {cls.YELLOW}>{cls.END} {msg}: ").strip()
    
    @classmethod
    def prompt_file(cls, msg):
        while True:
            path = cls.prompt(msg)
            if not path:
                cls.error("File path required.")
                continue
            if not os.path.exists(path):
                cls.error(f"File not found: {path}")
                continue
            return path
    
    @classmethod
    def confirm(cls, msg, default=True):
        hint = "Y/n" if default else "y/N"
        result = cls.prompt(f"{msg} ({hint})", "").lower()
        if not result:
            return default
        return result in ('y', 'yes')
    
    @classmethod
    def progress(cls, current, total, label="Progress"):
        pct = int(current / total * 100)
        bar = '█' * (pct // 5) + '░' * (20 - pct // 5)
        print(f"\r  {label}: {bar} {pct}%", end='', flush=True)
        if current == total:
            print()
    
    @classmethod
    def wait(cls):
        input(f"\n  {cls.CYAN}Press Enter to continue...{cls.END}")


# =============================================================================
# IMAGE PROCESSOR
# =============================================================================

class ImageProcessor:
    """Handles televiewer image cropping."""
    
    def __init__(self):
        self.cfg = Config()
    
    def detect_depth(self, height):
        """Auto-detect max depth from image height."""
        return int((height - self.cfg.DEPTH_OFFSET) / self.cfg.PIXELS_PER_FOOT)
    
    def get_row(self, depth, with_label=False):
        """Convert depth to pixel row."""
        row = self.cfg.DEPTH_OFFSET + int(depth * self.cfg.PIXELS_PER_FOOT)
        if with_label:
            row -= self.cfg.LABEL_OFFSET
        return max(0, row)
    
    def crop_image(self, image_path, output_dir):
        """
        Crop televiewer image into depth intervals.
        
        Returns:
            tuple: (list of image paths, max depth)
        """
        img = Image.open(image_path)
        width, height = img.size
        basename = Path(image_path).stem
        
        max_depth = self.detect_depth(height)
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate intervals: 0.5-3, then 3-6, 6-9, etc.
        intervals = [(0.5, 3.0)]
        for start in range(3, max_depth, 3):
            end = min(start + 3, max_depth)
            intervals.append((start, float(end)))
            if end >= max_depth:
                break
        
        output_files = []
        
        for i, (start, end) in enumerate(intervals):
            Console.progress(i + 1, len(intervals), "Cropping")
            
            if i == 0:
                # First image includes horizontal axis header
                body_start = self.get_row(start)
                body_end = self.get_row(end)
                crop_height = self.cfg.HEADER_HEIGHT + (body_end - body_start)
                
                crop = Image.new('RGB', (self.cfg.RIGHT_CROP_X, crop_height), 'white')
                header = img.crop((0, 0, self.cfg.RIGHT_CROP_X, self.cfg.HEADER_HEIGHT))
                body = img.crop((0, body_start, self.cfg.RIGHT_CROP_X, body_end))
                crop.paste(header, (0, 0))
                crop.paste(body, (0, self.cfg.HEADER_HEIGHT))
            else:
                # Subsequent images start at depth label
                row_start = self.get_row(start, with_label=True)
                row_end = min(self.get_row(end, with_label=True), height)
                crop = img.crop((0, row_start, self.cfg.RIGHT_CROP_X, row_end))
            
            # Generate filename
            end_str = str(int(end)) if end == int(end) else str(end)
            filename = f"{basename}_{start}-{end_str}ft.jpg"
            filepath = os.path.join(output_dir, filename)
            crop.save(filepath, 'JPEG', quality=95)
            output_files.append(filepath)
        
        return output_files, max_depth


# =============================================================================
# DATABASE CLIENT
# =============================================================================

class Database:
    """Supabase database client."""
    
    def __init__(self, api_key):
        self.url = Config.SUPABASE_URL
        self.key = api_key
    
    def _headers(self):
        return {
            "apikey": self.key,
            "Authorization": f"Bearer {self.key}",
            "Content-Type": "application/json",
            "Prefer": "return=representation"
        }
    
    def test(self):
        """Test connection."""
        try:
            r = requests.get(f"{self.url}/rest/v1/", headers=self._headers(), timeout=10)
            return r.status_code == 200
        except:
            return False
    
    def get_log(self, vb_id):
        """Get log by boring ID."""
        r = requests.get(
            f"{self.url}/rest/v1/optical_televiewer_logs?vb_id_txt=eq.{vb_id}",
            headers=self._headers()
        )
        data = r.json() if r.status_code == 200 else []
        return data[0] if data else None
    
    def list_logs(self):
        """List all logs."""
        r = requests.get(
            f"{self.url}/rest/v1/optical_televiewer_logs?select=vb_id_txt,drill_date&order=vb_id_txt",
            headers=self._headers()
        )
        return r.json() if r.status_code == 200 else []
    
    def save_log(self, data):
        """Insert or update log."""
        # Check if exists
        existing = self.get_log(data['vb_id_txt'])
        if existing:
            r = requests.patch(
                f"{self.url}/rest/v1/optical_televiewer_logs?vb_id_txt=eq.{data['vb_id_txt']}",
                headers=self._headers(),
                json=data
            )
        else:
            r = requests.post(
                f"{self.url}/rest/v1/optical_televiewer_logs",
                headers=self._headers(),
                json=data
            )
        return r.status_code in (200, 201)


# =============================================================================
# REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    """Generates Word reports from template."""
    
    def __init__(self):
        self.cfg = Config()
    
    def format_date(self, date_str):
        """Format date as 'Month DD, YYYY'."""
        if not date_str:
            return ""
        try:
            dt = datetime.strptime(str(date_str)[:10], "%Y-%m-%d")
            return dt.strftime("%B %d, %Y")
        except:
            return str(date_str)
    
    def find_image_cells(self, doc):
        """Find image area table indices."""
        indices = []
        for i, table in enumerate(doc.tables):
            if len(table.rows) == 1 and len(table.columns) == 1:
                if not table.rows[0].cells[0].text.strip():
                    indices.append(i)
        return indices
    
    def generate(self, template_path, images, header_data, output_path):
        """
        Generate the photo report.
        
        Args:
            template_path: Path to Word template
            images: List of cropped image paths
            header_data: Dictionary with header field values
            output_path: Output file path
        """
        doc = Document(template_path)
        image_indices = self.find_image_cells(doc)
        
        # Build replacements
        replacements = {
            'vb_id_txt': header_data.get('vb_id_txt', ''),
            'north_txt': header_data.get('north_txt', ''),
            'easti_txt': header_data.get('easti_txt', ''),
            'stat_num': header_data.get('stat_num', ''),
            'ground_elev_num': str(header_data.get('ground_elev_num', '')),
            'column_panel_txt': header_data.get('column_panel_txt', ''),
            'column_panel_joint_txt': header_data.get('column_panel_joint_txt', ''),
            'ct_txt': header_data.get('ct_txt', ''),
            'drill_date_date': self.format_date(header_data.get('drill_date')),
            'drill_by_txt': header_data.get('drill_by_txt', ''),
            'op_tv_logger': header_data.get('op_tv_logger', ''),
            'op_tv_date': self.format_date(header_data.get('op_tv_date'))
        }
        
        total_steps = len(images) + 2
        step = 0
        
        # Insert images
        for idx, img_path in zip(image_indices[:len(images)], images):
            step += 1
            Console.progress(step, total_steps, "Building report")
            
            table = doc.tables[idx]
            cell = table.rows[0].cells[0]
            
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_path, height=Inches(self.cfg.IMAGE_HEIGHT_INCHES))
        
        # Replace placeholders
        step += 1
        Console.progress(step, total_steps, "Building report")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.runs:
                            text = para.text
                            changed = False
                            for key, val in replacements.items():
                                if key in text:
                                    text = text.replace(key, str(val) if val else "")
                                    changed = True
                            if changed:
                                for run in para.runs:
                                    run.text = ''
                                para.runs[0].text = text
        
        # Add notes to final page
        if len(image_indices) > len(images):
            notes_idx = image_indices[len(images)]
            table = doc.tables[notes_idx]
            cell = table.rows[0].cells[0]
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            notes_para = cell.add_paragraph()
            run = notes_para.add_run(self.cfg.NOTES)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        
        # Save
        step += 1
        Console.progress(step, total_steps, "Building report")
        
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        
        return output_path


# =============================================================================
# MAIN APPLICATION
# =============================================================================

class App:
    """Main application."""
    
    def __init__(self):
        self.db = None
        self.api_key = None
    
    def run(self):
        """Run the application."""
        # Disable colors on Windows if needed
        if os.name == 'nt':
            try:
                import ctypes
                ctypes.windll.kernel32.SetConsoleMode(
                    ctypes.windll.kernel32.GetStdHandle(-11), 7)
            except:
                Console.CYAN = Console.GREEN = Console.YELLOW = ''
                Console.RED = Console.BOLD = Console.END = ''
        
        self.main_menu()
    
    def main_menu(self):
        """Show main menu."""
        while True:
            Console.header()
            Console.section("MAIN MENU")
            print("""
    [1] Generate Photo Report (with database)
    [2] Generate Photo Report (manual entry)
    [3] Crop Images Only
    [4] View Database Logs
    [Q] Quit
""")
            choice = Console.prompt("Select option").upper()
            
            if choice == '1':
                self.generate_with_database()
            elif choice == '2':
                self.generate_manual()
            elif choice == '3':
                self.crop_only()
            elif choice == '4':
                self.view_logs()
            elif choice == 'Q':
                Console.info("Goodbye!")
                print()
                break
    
    def connect_db(self):
        """Ensure database connection."""
        if self.db:
            return True
        
        Console.section("DATABASE CONNECTION")
        self.api_key = Console.prompt("Supabase service_role key")
        
        if not self.api_key:
            Console.error("API key required.")
            return False
        
        self.db = Database(self.api_key)
        
        Console.info("Testing connection...")
        if not self.db.test():
            Console.error("Connection failed. Check your API key.")
            self.db = None
            return False
        
        Console.success("Connected to database!")
        return True
    
    def generate_with_database(self):
        """Generate report using database for header data."""
        Console.header()
        Console.section("GENERATE PHOTO REPORT (Database Mode)")
        
        if not self.connect_db():
            Console.wait()
            return
        
        # Get boring ID
        vb_id = Console.prompt("Verification Boring ID (e.g., H7-VB-04)")
        if not vb_id:
            Console.error("Boring ID required.")
            Console.wait()
            return
        
        # Fetch from database
        Console.info(f"Looking up {vb_id}...")
        header_data = self.db.get_log(vb_id)
        
        if not header_data:
            Console.warning(f"No data found for {vb_id}")
            if Console.confirm("Enter header data manually?"):
                header_data = self.get_header_input(vb_id)
                if Console.confirm("Save to database?"):
                    if self.db.save_log(header_data):
                        Console.success("Saved to database!")
                    else:
                        Console.error("Failed to save.")
            else:
                Console.wait()
                return
        else:
            Console.success(f"Found: {header_data['vb_id_txt']}")
        
        # Get files
        self.generate_report(header_data)
    
    def generate_manual(self):
        """Generate report with manual header entry."""
        Console.header()
        Console.section("GENERATE PHOTO REPORT (Manual Entry)")
        
        vb_id = Console.prompt("Verification Boring ID")
        if not vb_id:
            Console.error("Boring ID required.")
            Console.wait()
            return
        
        header_data = self.get_header_input(vb_id)
        self.generate_report(header_data)
    
    def get_header_input(self, vb_id):
        """Get header data from user input."""
        Console.section("ENTER HEADER DATA")
        
        return {
            'vb_id_txt': vb_id,
            'north_txt': Console.prompt("Northing"),
            'easti_txt': Console.prompt("Easting"),
            'stat_num': Console.prompt("Station"),
            'ground_elev_num': Console.prompt("Ground Elevation", "0"),
            'column_panel_txt': Console.prompt("Column/Panel"),
            'column_panel_joint_txt': Console.prompt("Column/Panel Joint"),
            'ct_txt': Console.prompt("CT Number"),
            'drill_date': Console.prompt("Drill Date (YYYY-MM-DD)"),
            'drill_by_txt': Console.prompt("Drilled By"),
            'op_tv_logger': Console.prompt("TV Logger"),
            'op_tv_date': Console.prompt("TV Date (YYYY-MM-DD)")
        }
    
    def generate_report(self, header_data):
        """Generate the photo report."""
        Console.section("FILE SELECTION")
        
        # Get source image
        image_path = Console.prompt_file("Televiewer image file (JPG)")
        
        # Get template
        template_default = "Template_18_Pages_No_Images_Rev_03.docx"
        if os.path.exists(template_default):
            template_path = Console.prompt(f"Word template", template_default)
            if not template_path:
                template_path = template_default
        else:
            template_path = Console.prompt_file("Word template file")
        
        # Output directory
        output_dir = Console.prompt("Output directory", "./output")
        os.makedirs(output_dir, exist_ok=True)
        
        # Process
        Console.section("GENERATING REPORT")
        
        try:
            # Crop images
            Console.info("Cropping televiewer image...")
            processor = ImageProcessor()
            cropped_dir = os.path.join(output_dir, "cropped_images")
            images, max_depth = processor.crop_image(image_path, cropped_dir)
            Console.success(f"Created {len(images)} images (depth: {max_depth} ft)")
            
            # Generate report
            Console.info("Building Word report...")
            generator = ReportGenerator()
            vb_id = header_data['vb_id_txt']
            output_path = os.path.join(output_dir, f"Optical_Televiewer_Report_{vb_id}.docx")
            generator.generate(template_path, images, header_data, output_path)
            
            Console.success("Report generated successfully!")
            print(f"\n  Output: {output_path}")
            print(f"  Images: {cropped_dir}")
            
        except Exception as e:
            Console.error(f"Failed: {str(e)}")
        
        Console.wait()
    
    def crop_only(self):
        """Crop images without generating report."""
        Console.header()
        Console.section("CROP IMAGES ONLY")
        
        image_path = Console.prompt_file("Televiewer image file (JPG)")
        output_dir = Console.prompt("Output directory", "./cropped_images")
        
        Console.section("PROCESSING")
        
        try:
            processor = ImageProcessor()
            images, max_depth = processor.crop_image(image_path, output_dir)
            Console.success(f"Created {len(images)} images (depth: {max_depth} ft)")
            print(f"\n  Output: {output_dir}")
        except Exception as e:
            Console.error(f"Failed: {str(e)}")
        
        Console.wait()
    
    def view_logs(self):
        """View database logs."""
        Console.header()
        Console.section("DATABASE LOGS")
        
        if not self.connect_db():
            Console.wait()
            return
        
        logs = self.db.list_logs()
        
        if not logs:
            Console.warning("No logs found in database.")
        else:
            print(f"\n  {'Boring ID':<15} {'Drill Date':<15}")
            print(f"  {'-'*30}")
            for log in logs:
                drill = str(log.get('drill_date', ''))[:10]
                print(f"  {log['vb_id_txt']:<15} {drill:<15}")
            print(f"\n  Total: {len(logs)} logs")
        
        Console.wait()


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    try:
        App().run()
    except KeyboardInterrupt:
        print("\n\n  Interrupted. Goodbye!")
        sys.exit(0)
