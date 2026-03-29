#!/usr/bin/env python3
"""
Optical Televiewer Image Log Report Generator
==============================================

This script automates the generation of Optical Televiewer Image Log reports by:
1. Cropping downhole camera log images into 3-foot depth intervals
2. Pulling header data from a Supabase PostgreSQL database
3. Populating a Word template with images and data

Features:
- Auto-detect maximum depth from image dimensions
- Support for logs up to 66+ feet depth
- Preserve original scale ruler from source images
- Include horizontal axis header (azimuth labels) on first image
- Format dates as "Month DD, YYYY"
- Add notes to final page

Required Libraries:
    pip install Pillow python-docx requests

Required Files:
    - Word template: Template_18_Pages_No_Images_Rev_03.docx
    - Source image: Full-depth downhole log JPG

Supabase Database Table Schema:
    CREATE TABLE optical_televiewer_logs (
        id SERIAL PRIMARY KEY,
        vb_id_txt VARCHAR(15) NOT NULL,
        north_txt VARCHAR(15),
        easti_txt VARCHAR(15),
        stat_num VARCHAR(10),
        ground_elev_num NUMERIC(6,1),
        column_panel_txt VARCHAR(30),
        column_panel_joint_txt VARCHAR(30),
        ct_txt VARCHAR(10),
        drill_date DATE,
        drill_by_txt VARCHAR(15),
        op_tv_logger VARCHAR(25),
        op_tv_date DATE,
        image_file_path TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );

Author: Factory AI Droid
Version: 1.0.0
Date: March 2026
Repository: https://github.com/[your-username]/optical-televiewer-report-generator
"""

import os
import argparse
import requests
import json
from datetime import datetime
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """Configuration settings for the report generator."""
    
    # Supabase Configuration
    SUPABASE_URL = "https://oebbsdcdnnzqwdoacyyz.supabase.co"
    
    # Image Processing Parameters (calibrated from sample logs)
    PIXELS_PER_FOOT = 360
    DEPTH_OFFSET = -74          # Row where 0 feet would be
    HEADER_HEIGHT = 110         # Pixels - includes horizontal axis with azimuth labels
    RIGHT_CROP_X = 685          # Pixels - crop width (up to second 0° azimuth mark)
    LABEL_OFFSET = 20           # Pixels - labels appear above tick marks
    
    # Report Parameters
    IMAGE_HEIGHT_INCHES = 7.0
    MAX_SUPPORTED_DEPTH = 66    # feet
    
    # Notes text for final page
    NOTES_TEXT = """Notes:
(1) Camera type and recording settings are as follows:
Type: Digital Optical Borehole Televiewer (Mount Sopris)
Frame Rate: 20 HZ
Exposure: Medium
Lighting: 20
Resolution: 900
Logging Speed: 2.7-2.9 ft/min
(2) Photograph Not-to-Scale
(3) Borehole orientation presented in the above image is relative to the local magnetic field"""


# =============================================================================
# IMAGE CROPPING FUNCTIONS
# =============================================================================

class ImageCropper:
    """Handles cropping of downhole log images into depth intervals."""
    
    def __init__(self, config=None):
        """
        Initialize the cropper with configuration.
        
        Args:
            config: Configuration object (uses default Config if None)
        """
        self.config = config or Config()
    
    def get_row_for_depth(self, depth, include_label=False):
        """
        Calculate pixel row for a given depth.
        
        Args:
            depth: Depth in feet
            include_label: If True, start at top of label (20 pixels above tick)
        
        Returns:
            Pixel row number
        """
        row = self.config.DEPTH_OFFSET + int(depth * self.config.PIXELS_PER_FOOT)
        if include_label:
            row -= self.config.LABEL_OFFSET
        return max(0, row)
    
    def detect_max_depth(self, image_height):
        """
        Auto-detect maximum depth from image dimensions.
        
        Args:
            image_height: Image height in pixels
        
        Returns:
            Maximum depth in feet (integer)
        """
        max_depth = (image_height - self.config.DEPTH_OFFSET) / self.config.PIXELS_PER_FOOT
        return int(max_depth)
    
    def generate_intervals(self, max_depth):
        """
        Generate cropping intervals for a given max depth.
        
        Args:
            max_depth: Maximum depth in feet
        
        Returns:
            List of (start_depth, end_depth) tuples
        """
        # First interval: 0.5 to 3.0 ft (includes header)
        intervals = [(0.5, 3.0)]
        
        # Subsequent intervals: 3ft each
        for start in range(3, max_depth, 3):
            end = min(start + 3, max_depth)
            intervals.append((start, float(end)))
            if end >= max_depth:
                break
        
        return intervals
    
    def crop_image(self, image_path, output_dir, max_depth=None):
        """
        Crop a downhole log image into 3-foot intervals.
        
        Args:
            image_path: Path to the input image
            output_dir: Directory for output images
            max_depth: Maximum depth in feet (auto-detected if None)
        
        Returns:
            List of output file paths
        """
        # Load image
        img = Image.open(image_path)
        width, height = img.size
        basename = os.path.splitext(os.path.basename(image_path))[0]
        
        print(f"Processing: {image_path}")
        print(f"Image dimensions: {width}x{height} pixels")
        
        # Auto-detect max depth if not provided
        if max_depth is None:
            max_depth = self.detect_max_depth(height)
        
        print(f"Max depth: {max_depth} feet")
        
        # Create output directory if needed
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate intervals
        intervals = self.generate_intervals(max_depth)
        print(f"Generating {len(intervals)} images...")
        
        output_files = []
        
        for i, (start_depth, end_depth) in enumerate(intervals):
            if i == 0:
                # First image: include horizontal axis header + body from 0.5 to 3ft
                body_start = self.get_row_for_depth(start_depth)
                body_end = self.get_row_for_depth(end_depth)
                
                crop = Image.new('RGB', (self.config.RIGHT_CROP_X, 
                                         self.config.HEADER_HEIGHT + (body_end - body_start)), 'white')
                header = img.crop((0, 0, self.config.RIGHT_CROP_X, self.config.HEADER_HEIGHT))
                crop.paste(header, (0, 0))
                body = img.crop((0, body_start, self.config.RIGHT_CROP_X, body_end))
                crop.paste(body, (0, self.config.HEADER_HEIGHT))
            else:
                # Other images: start at the TOP of the depth label
                row_start = self.get_row_for_depth(start_depth, include_label=True)
                row_end = self.get_row_for_depth(end_depth, include_label=True)
                row_end = min(row_end, height)  # Don't exceed image bounds
                crop = img.crop((0, row_start, self.config.RIGHT_CROP_X, row_end))
            
            # Format filename
            if end_depth == int(end_depth):
                filename = f"{basename}_{start_depth}-{int(end_depth)}ft.jpg"
            else:
                filename = f"{basename}_{start_depth}-{end_depth}ft.jpg"
            filename = filename.replace('.0-', '-').replace('0.5', '0.5')
            
            output_path = os.path.join(output_dir, filename)
            crop.save(output_path, 'JPEG', quality=95)
            output_files.append(output_path)
            
            print(f"  {i+1}. {filename}: {crop.size}")
        
        return output_files


# =============================================================================
# DATABASE FUNCTIONS
# =============================================================================

class SupabaseClient:
    """Handles communication with Supabase PostgreSQL database."""
    
    def __init__(self, url=None, service_role_key=None):
        """
        Initialize the Supabase client.
        
        Args:
            url: Supabase project URL
            service_role_key: Supabase service role key
        """
        self.url = url or Config.SUPABASE_URL
        self.service_role_key = service_role_key
    
    def _get_headers(self):
        """Get headers for API requests."""
        if not self.service_role_key:
            raise ValueError("Supabase service role key not provided")
        
        return {
            "apikey": self.service_role_key,
            "Authorization": f"Bearer {self.service_role_key}",
            "Content-Type": "application/json",
            "Prefer": "return=representation"
        }
    
    def query_log(self, vb_id):
        """
        Query the database for log header data.
        
        Args:
            vb_id: Verification Boring ID (e.g., "H7-VB-04")
        
        Returns:
            Dictionary with log data, or None if not found
        """
        response = requests.get(
            f"{self.url}/rest/v1/optical_televiewer_logs?vb_id_txt=eq.{vb_id}",
            headers=self._get_headers()
        )
        
        if response.status_code == 200 and response.json():
            return response.json()[0]
        return None
    
    def insert_log(self, log_data):
        """
        Insert new log data into the database.
        
        Args:
            log_data: Dictionary with log field values
        
        Returns:
            Inserted row data, or None if failed
        """
        response = requests.post(
            f"{self.url}/rest/v1/optical_televiewer_logs",
            headers=self._get_headers(),
            json=log_data
        )
        
        if response.status_code == 201:
            return response.json()
        return None
    
    def update_log(self, vb_id, updates):
        """
        Update existing log data in the database.
        
        Args:
            vb_id: Verification Boring ID
            updates: Dictionary with fields to update
        
        Returns:
            Updated row data, or None if failed
        """
        response = requests.patch(
            f"{self.url}/rest/v1/optical_televiewer_logs?vb_id_txt=eq.{vb_id}",
            headers=self._get_headers(),
            json=updates
        )
        
        if response.status_code == 200:
            return response.json()
        return None
    
    def list_logs(self):
        """
        List all logs in the database.
        
        Returns:
            List of log records
        """
        response = requests.get(
            f"{self.url}/rest/v1/optical_televiewer_logs?select=id,vb_id_txt,created_at&order=created_at.desc",
            headers=self._get_headers()
        )
        
        if response.status_code == 200:
            return response.json()
        return []


# =============================================================================
# REPORT GENERATION FUNCTIONS
# =============================================================================

class ReportGenerator:
    """Generates Word reports from templates and data."""
    
    def __init__(self, config=None):
        """
        Initialize the report generator.
        
        Args:
            config: Configuration object (uses default Config if None)
        """
        self.config = config or Config()
    
    @staticmethod
    def format_date(date_str):
        """
        Convert YYYY-MM-DD to 'Month DD, YYYY' format.
        
        Args:
            date_str: Date string in YYYY-MM-DD format
        
        Returns:
            Formatted date string (e.g., "February 25, 2026")
        """
        if date_str:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            return dt.strftime("%B %d, %Y")
        return ""
    
    def build_replacements(self, log_data):
        """
        Build replacement dictionary from log data.
        
        Args:
            log_data: Dictionary from database query
        
        Returns:
            Dictionary mapping placeholder names to values
        """
        return {
            'vb_id_txt': log_data.get('vb_id_txt', ''),
            'north_txt': log_data.get('north_txt', ''),
            'easti_txt': log_data.get('easti_txt', ''),
            'stat_num': log_data.get('stat_num', ''),
            'ground_elev_num': str(log_data.get('ground_elev_num', '')),
            'column_panel_txt': log_data.get('column_panel_txt', ''),
            'column_panel_joint_txt': log_data.get('column_panel_joint_txt', ''),
            'ct_txt': log_data.get('ct_txt', ''),
            'drill_date_date': self.format_date(log_data.get('drill_date')),
            'drill_by_txt': log_data.get('drill_by_txt', ''),
            'op_tv_logger': log_data.get('op_tv_logger', ''),
            'op_tv_date': self.format_date(log_data.get('op_tv_date'))
        }
    
    def find_image_tables(self, doc):
        """
        Find indices of image area tables in the template.
        
        Args:
            doc: python-docx Document object
        
        Returns:
            List of table indices for image areas
        """
        image_table_indices = []
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            if rows == 1 and cols == 1:
                first_text = table.rows[0].cells[0].text.strip()
                if not first_text:
                    image_table_indices.append(i)
        return image_table_indices
    
    def insert_images(self, doc, image_files, image_table_indices):
        """
        Insert images into the Word document.
        
        Args:
            doc: python-docx Document object
            image_files: List of image file paths
            image_table_indices: List of table indices for image areas
        """
        for page_num, (table_idx, img_path) in enumerate(
                zip(image_table_indices[:len(image_files)], image_files)):
            table = doc.tables[table_idx]
            cell = table.rows[0].cells[0]
            
            # Clear existing content
            for para in cell.paragraphs:
                p_elem = para._element
                p_elem.getparent().remove(p_elem)
            
            # Add centered image
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, height=Inches(self.config.IMAGE_HEIGHT_INCHES))
    
    def replace_placeholders(self, doc, replacements):
        """
        Replace all placeholder variables in the document.
        
        Args:
            doc: python-docx Document object
            replacements: Dictionary mapping placeholder names to values
        
        Returns:
            Number of replacements made
        """
        replacement_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            full_text = paragraph.text
                            modified = False
                            for var_name, value in replacements.items():
                                if var_name in full_text:
                                    full_text = full_text.replace(
                                        var_name, str(value) if value else "")
                                    modified = True
                                    replacement_count += 1
                            
                            if modified:
                                for run in paragraph.runs:
                                    run.text = ''
                                paragraph.runs[0].text = full_text
        
        return replacement_count
    
    def add_notes(self, doc, image_table_indices, notes_page_index=17):
        """
        Add notes to the final page of the report.
        
        Args:
            doc: python-docx Document object
            image_table_indices: List of table indices for image areas
            notes_page_index: Index of the page for notes (0-based)
        """
        if len(image_table_indices) > notes_page_index:
            last_table_idx = image_table_indices[notes_page_index]
            last_table = doc.tables[last_table_idx]
            cell = last_table.rows[0].cells[0]
            
            # Clear existing content
            for para in cell.paragraphs:
                p_elem = para._element
                p_elem.getparent().remove(p_elem)
            
            # Add notes
            notes_para = cell.add_paragraph()
            notes_run = notes_para.add_run(self.config.NOTES_TEXT)
            notes_run.font.size = Pt(11)
            notes_run.font.name = 'Arial'
    
    def generate(self, template_path, image_files, log_data, output_path):
        """
        Generate a complete report.
        
        Args:
            template_path: Path to the Word template
            image_files: List of cropped image file paths
            log_data: Dictionary with log header data
            output_path: Path for the output report
        
        Returns:
            Path to the generated report
        """
        # Load template
        doc = Document(template_path)
        print(f"Template loaded: {len(doc.tables)} tables")
        
        # Find image tables
        image_table_indices = self.find_image_tables(doc)
        print(f"Found {len(image_table_indices)} image areas")
        
        # Insert images
        self.insert_images(doc, image_files, image_table_indices)
        print(f"Inserted {len(image_files)} images")
        
        # Replace placeholders
        replacements = self.build_replacements(log_data)
        count = self.replace_placeholders(doc, replacements)
        print(f"Made {count} replacements")
        
        # Add notes to final page
        self.add_notes(doc, image_table_indices)
        print("Notes added to final page")
        
        # Save report
        doc.save(output_path)
        print(f"Saved: {output_path}")
        
        return output_path


# =============================================================================
# MAIN WORKFLOW FUNCTIONS
# =============================================================================

def generate_report(vb_id, source_image_path, template_path, output_dir,
                    service_role_key, output_filename=None):
    """
    Generate a complete Optical Televiewer report.
    
    This is the main entry point for the complete workflow.
    
    Args:
        vb_id: Verification Boring ID (e.g., "H7-VB-04")
        source_image_path: Path to the full-depth source image
        template_path: Path to the Word template
        output_dir: Directory for output files
        service_role_key: Supabase service role key
        output_filename: Custom output filename (auto-generated if None)
    
    Returns:
        Path to the generated report
    """
    print("=" * 60)
    print("OPTICAL TELEVIEWER REPORT GENERATION")
    print("=" * 60)
    
    # Initialize components
    db = SupabaseClient(service_role_key=service_role_key)
    cropper = ImageCropper()
    generator = ReportGenerator()
    
    # Step 1: Query database
    print(f"\n[STEP 1] Querying database for: {vb_id}")
    log_data = db.query_log(vb_id)
    if not log_data:
        raise ValueError(f"Log data not found for: {vb_id}")
    print(f"   Found: {log_data['vb_id_txt']}")
    
    # Step 2: Crop source image
    print(f"\n[STEP 2] Cropping source image")
    cropped_dir = os.path.join(output_dir, "cropped_images")
    image_files = cropper.crop_image(source_image_path, cropped_dir)
    print(f"   Generated {len(image_files)} images")
    
    # Step 3: Generate report
    print(f"\n[STEP 3] Generating report")
    if output_filename is None:
        output_filename = f"Optical_Televiewer_Report_{vb_id}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    generator.generate(template_path, image_files, log_data, output_path)
    
    print("\n" + "=" * 60)
    print("COMPLETE!")
    print("=" * 60)
    print(f"\nReport saved to: {output_path}")
    
    return output_path


def crop_images_only(source_image_path, output_dir, max_depth=None):
    """
    Crop a source image without generating a report.
    
    Useful for processing images before database entry is created.
    
    Args:
        source_image_path: Path to the full-depth source image
        output_dir: Directory for output images
        max_depth: Maximum depth in feet (auto-detected if None)
    
    Returns:
        List of output file paths
    """
    cropper = ImageCropper()
    return cropper.crop_image(source_image_path, output_dir, max_depth)


def list_database_logs(service_role_key):
    """
    List all logs in the database.
    
    Args:
        service_role_key: Supabase service role key
    
    Returns:
        List of log records
    """
    db = SupabaseClient(service_role_key=service_role_key)
    return db.list_logs()


def add_log_to_database(log_data, service_role_key):
    """
    Add a new log entry to the database.
    
    Args:
        log_data: Dictionary with log field values
        service_role_key: Supabase service role key
    
    Returns:
        Inserted row data
    """
    db = SupabaseClient(service_role_key=service_role_key)
    return db.insert_log(log_data)


# =============================================================================
# COMMAND LINE INTERFACE
# =============================================================================

def main():
    """Main entry point for command-line usage."""
    parser = argparse.ArgumentParser(
        description="Optical Televiewer Image Log Report Generator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Generate a full report:
    python %(prog)s generate --vb-id H7-VB-04 --image log.jpg \\
        --template template.docx --output-dir ./output --api-key YOUR_KEY

  Crop images only:
    python %(prog)s crop --image log.jpg --output-dir ./cropped

  List database logs:
    python %(prog)s list --api-key YOUR_KEY
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Command to run')
    
    # Generate report command
    gen_parser = subparsers.add_parser('generate', help='Generate a complete report')
    gen_parser.add_argument('--vb-id', required=True, 
                            help='Verification Boring ID (e.g., H7-VB-04)')
    gen_parser.add_argument('--image', required=True, 
                            help='Path to source image')
    gen_parser.add_argument('--template', required=True, 
                            help='Path to Word template')
    gen_parser.add_argument('--output-dir', default='./output', 
                            help='Output directory')
    gen_parser.add_argument('--api-key', required=True, 
                            help='Supabase service role key')
    gen_parser.add_argument('--output-name', 
                            help='Custom output filename')
    
    # Crop images command
    crop_parser = subparsers.add_parser('crop', help='Crop images only')
    crop_parser.add_argument('--image', required=True, 
                             help='Path to source image')
    crop_parser.add_argument('--output-dir', default='./cropped', 
                             help='Output directory')
    crop_parser.add_argument('--max-depth', type=int, 
                             help='Maximum depth in feet (auto-detected if not specified)')
    
    # List logs command
    list_parser = subparsers.add_parser('list', help='List database logs')
    list_parser.add_argument('--api-key', required=True, 
                             help='Supabase service role key')
    
    args = parser.parse_args()
    
    if args.command == 'generate':
        generate_report(
            vb_id=args.vb_id,
            source_image_path=args.image,
            template_path=args.template,
            output_dir=args.output_dir,
            service_role_key=args.api_key,
            output_filename=args.output_name
        )
    
    elif args.command == 'crop':
        files = crop_images_only(
            source_image_path=args.image,
            output_dir=args.output_dir,
            max_depth=args.max_depth
        )
        print(f"\nGenerated {len(files)} cropped images")
    
    elif args.command == 'list':
        logs = list_database_logs(args.api_key)
        print(f"\nFound {len(logs)} logs in database:")
        for log in logs:
            print(f"  {log['id']}: {log['vb_id_txt']} (created: {log['created_at']})")
    
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
